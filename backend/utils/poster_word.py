"""
Generate Word documents for posters.
Page 1: Metadata (title, authors, affiliations, abstract)
Page 2: Poster image (embedded)
"""

import base64
import io
from typing import Dict, Any
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def generate_poster_word(poster: Dict[str, Any]) -> bytes:
    """
    Generate Word document for poster.

    Args:
        poster: Dict with keys:
            - title: str
            - authors: list of {"first_name": str, "last_name": str, "affiliation": str}
            - abstract: str
            - poster_image: str (base64 encoded image data)

    Returns:
        Word document as bytes
    """
    doc = Document()

    # Set up document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # PAGE 1: METADATA
    # Title
    title = poster.get("title", "Untitled Poster")
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(title)
    title_run.font.size = Pt(28)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(102, 102, 153)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.space_after = Pt(15)

    # Authors
    authors = poster.get("authors", [])
    if authors:
        for author in authors:
            first_name = author.get("first_name", "")
            last_name = author.get("last_name", "")
            affiliation = author.get("affiliation", "")

            name = f"{first_name} {last_name}".strip()
            if name:
                author_para = doc.add_paragraph(name)
                author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                author_para.runs[0].font.size = Pt(12)
                author_para.runs[0].font.bold = True

                if affiliation:
                    aff_para = doc.add_paragraph(affiliation)
                    aff_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    aff_para.runs[0].font.size = Pt(10)
                    aff_para.runs[0].font.italic = True

        doc.add_paragraph()  # Spacing

    # Abstract
    abstract = poster.get("abstract", "")
    if abstract:
        abstract_heading = doc.add_paragraph()
        abstract_heading_run = abstract_heading.add_run("Abstract")
        abstract_heading_run.font.bold = True
        abstract_heading_run.font.size = Pt(12)
        abstract_heading.space_after = Pt(6)

        abstract_para = doc.add_paragraph(abstract)
        abstract_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for run in abstract_para.runs:
            run.font.size = Pt(11)

    # Page break
    doc.add_page_break()

    # PAGE 2: POSTER IMAGE
    poster_image = poster.get("poster_image", "")
    if poster_image:
        try:
            # Decode base64 image
            if "," in poster_image:
                # Data URI format (data:image/png;base64,...)
                poster_image = poster_image.split(",")[1]

            image_bytes = base64.b64decode(poster_image)
            image_stream = io.BytesIO(image_bytes)

            # Add image
            poster_para = doc.add_paragraph()
            poster_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = poster_para.add_run()
            run.add_picture(image_stream, width=Inches(6))

        except Exception as e:
            # If image fails, add error message
            error_para = doc.add_paragraph(f"[Image could not be embedded: {str(e)}]")
            error_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Convert to bytes
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output.read()
