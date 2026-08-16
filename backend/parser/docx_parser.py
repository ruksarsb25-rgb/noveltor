"""
DOCX parsing logic using python-docx.
Extracts structured metadata and content from academic manuscripts.
"""
import re
import html as _html_lib
from docx import Document
from docx.oxml.ns import qn as _qn
from docx.text.hyperlink import Hyperlink
from docx.text.run import Run
from lxml import etree as _etree
from utils.equations import extract_equation_text, mathml_from_omml, omml_to_latex
from utils.latex_to_mathml import detect_latex_formulas, latex_to_mathml

# Translate Unicode superscript digits → ASCII digits
_SUP_TO_NUM = str.maketrans('¹²³⁴⁵⁶⁷⁸⁹⁰', '1234567890')

_EMAIL_RE = re.compile(r'[\w.+-]+@[\w.-]+\.\w+')
_CORRESP_LINE_RE = re.compile(r'\*?\s*corresponding\s+authors?\s*[:\-]', re.IGNORECASE)

# Matches Unicode superscript runs OR digits glued directly to a letter (regular superscripts)
# Updated to also match asterisks before digits (like N*1, N*2) and after (like R3*, R4*)
# (?:,\d+)* / (?:,[¹²³⁴⁵⁶⁷⁸⁹⁰]+)* absorbs further comma-joined numbers so an author
# affiliated with more than one institution — "Chan5,6" — is captured as ONE marker
# ("5,6"). Without it, \d+ only ever grabs the first digit run; the second number
# (and its leading comma) falls outside the match and gets glued onto the START of
# the next author's name instead (e.g. "6, H. C. Ananda Murthy" as one "name").
_AUTHOR_MARKER_RE = re.compile(
    r'([¹²³⁴⁵⁶⁷⁸⁹⁰]+(?:,[¹²³⁴⁵⁶⁷⁸⁹⁰]+)*,?\*?|(?<=[A-Za-z])\*?\d+(?:,\d+)*,?\*?)'
)

# An affiliation line prefixed with its number — "1Department of...",
# "2 Department of..." (some authors leave a stray space after the digit).
# One-or-two-digit affiliation number directly followed by a letter (with at
# most a single space between): deliberately excludes "2.1 " so a genuine
# numbered subsection heading is never mistaken for an affiliation line.
_AFFIL_LINE_RE = re.compile(r'^(?:[¹²³⁴⁵⁶⁷⁸⁹⁰]+|\d{1,2})\s?[A-Za-z]')

# Heading / structure detection
# Numbered heading: "1. Introduction", "2.1 Chemicals", "3.2.1 Sub" — must start capital after number
_NUMBERED_HEADING_RE = re.compile(r'^(\d+)([\.\d]*?)\.?\s+([A-Z])')
_ABSTRACT_RE = re.compile(r'^abstract\b', re.IGNORECASE)
_KEYWORDS_RE = re.compile(r'^key\s*words?\s*[:\-]', re.IGNORECASE)
# "References", "5. References", "5 References" — at end of heading
# Matches: "References", "References:", "References;", "References :", "References .",
#           "5. References", "5 References:", etc.
_REFS_HEADING_RE = re.compile(r'^(?:\d+[\.\d]*\.?\s+)?references?\s*[;:.,]?\s*$', re.IGNORECASE)
# Reference list entry: "[1] ..." or "1. ..." — numbered formats
_REF_ENTRY_RE = re.compile(r'^\[?\d+[\]\.]\s+\S')
# Bullet characters that start a reference item
_REF_BULLET_RE = re.compile(r'^[•\-\*●◆▪]\s+')
# DOI in any format: "doi: 10.xxx", "https://doi.org/10.xxx", or bare "10.xxxx/yyyy"
_DOI_RE = re.compile(r'(?:doi:\s*|https?://doi\.org/)?([0-9]{2}\.[0-9]{4}/[^\s,;<>\)]+)', re.IGNORECASE)
# Matches a DOI inside a hyperlink target, e.g. https://doi.org/10.1007/978-981-96-6795-6_18
_DOI_URL_RE = re.compile(r'doi\.org/(10\.\d{4,}/\S+)', re.IGNORECASE)

# Article-type label line ("Research paper", "Review Article", ...) that
# many journal templates place on its own line directly above the real
# title. Without this, the pre_title fallback rule (below, "still no title
# after N paragraphs") mistakes the label itself for the title, and the
# real title — appearing right after — gets swallowed into author-name
# parsing instead (phase has already moved past pre_title by then).
_ARTICLE_TYPE_LABELS = {
    'research paper', 'research article', 'review article', 'review paper',
    'original article', 'original research', 'original research article',
    'short communication', 'case report', 'case study', 'technical note',
    'mini review', 'mini-review', 'letter', 'communication', 'perspective',
    'editorial', 'commentary', 'brief report', 'rapid communication',
    'conference paper', 'conference proceeding',
}

_W_TBL     = _qn('w:tbl')
_W_P       = _qn('w:p')
_W_R       = _qn('w:r')
_W_T       = _qn('w:t')
_W_VALIGN  = _qn('w:vertAlign')
_W_VAL     = _qn('w:val')
_W_RPR     = _qn('w:rPr')
_W_HYPERLINK = _qn('w:hyperlink')


def _run_text_with_fmt(run) -> str:
    """Escape a run's text and wrap it in <sub>/<sup> per its vertAlign formatting."""
    t = run.text
    if not t:
        return ""
    rpr = run._element.find(_W_RPR)
    vert = None
    if rpr is not None:
        va = rpr.find(_W_VALIGN)
        if va is not None:
            vert = va.get(_W_VAL)
    escaped = _html_lib.escape(t, quote=False)
    if vert == 'subscript':
        return f'<sub>{escaped}</sub>'
    elif vert == 'superscript':
        return f'<sup>{escaped}</sup>'
    return escaped


def _para_text_with_fmt(p) -> str:
    """Extract paragraph text in document order, including hyperlink display text
    (e.g. linked author names) so it isn't dropped. A hyperlink's target is appended
    only when it's a DOI link whose DOI isn't already visible in the display text —
    this preserves CrossRef DOIs that documents show as plain link text like "here"."""
    parts = []
    doi_appendix = []

    for item in p.iter_inner_content():
        if isinstance(item, Hyperlink):
            link_text = ''.join(_run_text_with_fmt(r) for r in item.runs)
            parts.append(link_text)
            url = item.address or ""
            m = _DOI_URL_RE.search(url)
            if m and m.group(1) not in link_text:
                doi_appendix.append(url)
        else:
            parts.append(_run_text_with_fmt(item))

    text = ''.join(parts)
    if doi_appendix:
        text += " " + " ".join(doi_appendix)
    return text


def _cell_text_with_fmt(cell) -> str:
    """Extract all text from a table cell across all paragraphs, preserving sub/sup and hyperlink text."""
    parts = []
    for para in cell.paragraphs:
        if parts:
            parts.append(' ')

        for item in para.iter_inner_content():
            if isinstance(item, Hyperlink):
                link_text = ''.join(_run_text_with_fmt(r) for r in item.runs)
                parts.append(link_text)
                url = item.address or ""
                m = _DOI_URL_RE.search(url)
                if m and m.group(1) not in link_text:
                    parts.append(' ' + url)
            else:
                parts.append(_run_text_with_fmt(item))

    return ''.join(parts).strip()

# Drawing XML namespaces for inline/anchor image detection
_DRAWING_NS = "{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}"
_DRAW_A_NS  = "{http://schemas.openxmlformats.org/drawingml/2006/main}"
_REL_NS     = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}"
# VML namespace (older DOCX files use v:imagedata instead of DrawingML)
_VML_NS     = "{urn:schemas-microsoft-com:vml}"
_O_NS       = "{urn:schemas-microsoft-com:office:office}"
# OMML namespace for Word equations
_MATH_NS    = "{http://schemas.openxmlformats.org/officeDocument/2006/math}"
# Word "Group" shapes (wpg:wgp/wpg:grpSp) — used when an author selects
# several pictures and groups them into one object, e.g. a data table
# image overlaid on a chart image to build one composite figure panel.
_WPG_NS     = "{http://schemas.microsoft.com/office/word/2010/wordprocessingGroup}"
_PIC_NS     = "{http://schemas.openxmlformats.org/drawingml/2006/picture}"

# Figure label: Fig/Figure, optional dot, then space/hyphen/nothing, optional parens around number
_FIG_CAPTION_RE    = re.compile(r'^fig(?:ure)?\.?[\s\-]?\s*\(?\d', re.IGNORECASE)
_FIG_CAPTION_FULL  = re.compile(r'^fig(?:ure)?\.?[\s\-]?\s*\(?(\d+)\)?', re.IGNORECASE)
# Table label: same separator rules as figures
_TABLE_CAPTION_RE  = re.compile(r'^table\.?[\s\-]?\s*\(?\d', re.IGNORECASE)
_TABLE_CAPTION_FULL = re.compile(r'^table\.?[\s\-]?\s*\(?(\d+)\)?', re.IGNORECASE)
# Non-numbered items that should never get a Fig-N tag
_SKIP_FIG_RE       = re.compile(r'^(?:graphical\s+abstract|sch(?:ema|eme)(?:[-\s]\d*)?)\b', re.IGNORECASE)
# Body-text figure references: "Fig. 2 presents/shows/illustrates/is/demonstrates..."
_FIG_REF_VERB_RE   = re.compile(
    r'^fig(?:ure)?\.?[\s\-]?\s*\(?(\d+)\)?[a-z]?\s+(?:present|show|illustrat|depict|display|is\b|are\b|demonstrat)',
    re.IGNORECASE
)


# Strip leading "Fig. N:" / "Figure (12)." / "Table 1 -" / "Fig. 1 " prefix from caption text
_FIG_LABEL_PREFIX_RE = re.compile(
    r'^(?:fig(?:ure)?|table)\.?[\s\-]?\s*\(?\d+\)?[a-z]?\s*[:\.\-\s]+',
    re.IGNORECASE
)


def _strip_fig_label(text: str) -> str:
    """Remove 'Fig. N:' / 'Figure (12).' prefix; keep only the description."""
    return _FIG_LABEL_PREFIX_RE.sub("", text).strip()


def _extract_latex_blocks(text: str) -> list:
    """
    Split text by LaTeX formulas and return list of content blocks.

    Returns: [{"type": "text", "text": "..."}, {"type": "equation", "latex": "...", "mathml": "..."}]
    """
    formulas = detect_latex_formulas(text)
    if not formulas:
        return [{"type": "text", "text": text}]

    blocks = []
    last_end = 0

    for start, end, formula_text, is_display in formulas:
        # Text before formula
        if start > last_end:
            before_text = text[last_end:start].strip()
            if before_text:
                blocks.append({"type": "text", "text": before_text})

        # Formula block
        mathml = latex_to_mathml(formula_text, display=is_display)
        blocks.append({
            "type": "equation",
            "text": formula_text,  # Original LaTeX for reference
            "latex": formula_text,
            "mathml": mathml,
            "data_uri": ""  # No image for LaTeX-based equations
        })

        last_end = end

    # Remaining text
    if last_end < len(text):
        remaining = text[last_end:].strip()
        if remaining:
            blocks.append({"type": "text", "text": remaining})

    return blocks


def _collect_fig_captions(doc) -> dict:
    """Pre-scan entire document and return {fig_number (int): caption_text}.

    Prefers true caption lines (e.g. 'Fig. 2. SEM micrographs...') over
    body-text references (e.g. 'Fig. 2 presents the SEM...').
    """
    captions: dict[int, tuple[int, str]] = {}  # num → (priority, text); priority 0=caption,1=ref

    for p in doc.paragraphs:
        text = p.text.strip()
        m = _FIG_CAPTION_FULL.match(text)
        if not m:
            continue
        num = int(m.group(1))
        is_ref = bool(_FIG_REF_VERB_RE.match(text))
        priority = 1 if is_ref else 0  # 0 = true caption (preferred)

        existing = captions.get(num)
        if existing is None:
            captions[num] = (priority, text)
        else:
            ex_pri, ex_text = existing
            # Prefer lower priority (real caption over body ref); break ties by length
            if priority < ex_pri or (priority == ex_pri and len(text) > len(ex_text)):
                captions[num] = (priority, text)

    return {num: text for num, (_, text) in captions.items()}


def parse_docx(file_path: str) -> dict:
    doc = Document(file_path)

    fig_captions = _collect_fig_captions(doc)

    state = {
        "title": "",
        "authors_raw": "",
        "abstract": "",
        "keywords": [],
        "sections": [],
        "references": [],
        "figures": [],
    }

    _extract_structure(doc, state, fig_captions)

    return {
        "title": state["title"],
        "authors": _parse_authors(state["authors_raw"]),
        "abstract": state["abstract"],
        "keywords": state["keywords"],
        "sections": state["sections"],
        "references": _parse_references(state["references"]),
        "figures": state["figures"],
        # Journal identity is set by the editor, never inferred from manuscript content
        "journal_name": None,
        "publisher_name": None,
        "publisher_loc": None,
    }


def _classify_heading(text: str, style: str, is_bold: bool = False) -> str | None:
    """
    Return 'h1', 'h2', 'h3', or None.

    Priority order:
    1. Numbered heading regex — most reliable for NFP manuscripts:
       "1. Introduction" → h2, "2.1 Chemicals" → h3
    2. DOCX style name — used when no number prefix is present.
    3. Bold fallback (only when is_bold=True) — for poster/unnumbered documents
       where section headings are bold but have no number prefix and no Heading
       style applied. A bold paragraph is treated as h2 when:
         • ≤ 10 words (headings are short labels, not full sentences)
         • starts with an uppercase letter
         • contains no mid-sentence period (rules out bold sentences in body text)
    """
    m = _NUMBERED_HEADING_RE.match(text)
    if m:
        sub_part = m.group(2)  # everything between first digit and the space
        # If there are digits after the first number (e.g. ".1", ".2.3"), it's a subsection
        has_sub = bool(re.search(r'\d', sub_part))
        return "h3" if has_sub else "h2"

    if style == "Heading 1":
        return "h1"
    if style == "Heading 2":
        return "h2"
    if style == "Heading 3":
        return "h3"

    # ── Bold fallback for unnumbered / poster documents ───────────────────────
    if is_bold and text:
        words = text.split()
        if (
            len(words) <= 10                        # short — heading-like
            and text[0].isupper()                   # starts with capital
            and not re.search(r'\.\s+[A-Z]', text) # no "sentence. Next" pattern
            and not text.endswith(".")              # not a full sentence
        ):
            return "h2"

    return None


def _extract_structure(doc, state: dict, fig_captions: dict = None):
    """
    State-machine pass over body elements in document order (paragraphs and tables).
    Phases: pre_title → authors → abstract → body → refs

    Sections use a content array of blocks instead of a plain body string, so that
    tables can be inserted inline at their exact position in the flow.
    """
    if fig_captions is None:
        fig_captions = {}
    phase = "pre_title"
    current_section = None
    authors_raw_lines = []

    body_children = list(doc.element.body)
    body_len  = len(body_children)
    para_idx  = 0
    tbl_idx   = 0
    fig_counter = 0
    last_nonempty_text = ""
    idx = 0

    while idx < body_len:
        el = body_children[idx]
        idx += 1
        # ── Table element ─────────────────────────────────────────────────────
        if el.tag == _W_TBL:
            tbl = doc.tables[tbl_idx]
            tbl_idx += 1
            if phase == "body":
                tbl_block = _table_to_block(tbl, tbl_idx)
                if tbl_block:
                    if current_section is None:
                        current_section = _new_section("", "Other")
                    if current_section["subsections"]:
                        current_section["subsections"][-1]["content"].append(tbl_block)
                    else:
                        current_section["content"].append(tbl_block)
            continue

        # ── Paragraph element ─────────────────────────────────────────────────
        if el.tag != _W_P:
            continue

        p = doc.paragraphs[para_idx]
        para_idx += 1

        style_name = p.style.name if p.style else ""
        text = p.text.strip()
        if text:
            last_nonempty_text = text
        is_bold = all(run.bold for run in p.runs if run.text.strip()) and bool(p.runs)
        is_list = "list" in style_name.lower()
        font_size = None
        for run in p.runs:
            if run.font.size:
                font_size = run.font.size.pt
                break

        # ── OMML equation detection (before blank-text skip) ─────────────────
        # Detect equations in ANY phase, but handle specially for pre-body
        if _has_math(p._element):
            combined_text, omml_list, has_prose = _split_para_math(p)

            if current_section is None:
                current_section = _new_section("", "Other")
            target = (current_section["subsections"][-1]["content"]
                      if current_section["subsections"] else current_section["content"])

            if has_prose or not omml_list:
                # Prose paragraph with inline math mid-sentence (e.g. "Where:
                # C0- Initial concentration ... and Ct- Concentration...") —
                # render as ordinary paragraph text so nothing outside the
                # equation is lost, with <sub>/<sup> preserved for variables.
                if combined_text:
                    target.append({"type": "paragraph", "text": combined_text})
            else:
                # Standalone display equation (e.g. "D = Kλ/(β cos θ) -----(1)")
                data_uri = _math_para_to_image(p)
                omml   = omml_list[0]
                mathml = mathml_from_omml(omml) if omml else ""
                eq_text = " ".join(extract_equation_text(o) for o in omml_list).strip()
                eq_latex = " ".join(omml_to_latex(o) for o in omml_list).strip()

                if omml or data_uri:
                    target.append({
                        "type": "equation",
                        "omml": omml if omml else None,
                        "mathml": mathml if mathml else None,
                        "text": eq_text if eq_text else None,  # Extracted text with proper subscripts
                        "latex": eq_latex if eq_latex else None,  # Real LaTeX math (\frac, \beta, …) for /export/pdf-latex
                        "data_uri": data_uri if data_uri else None
                    })
            continue

        # ── Inline image detection (before blank-text skip) ──────────────────
        if _has_image(p._element):
            caption = ""
            skip_label = skip_caption = ""
            idx_delta = para_delta = 0

            # Look back: preceding non-empty paragraph was a skip label
            is_skip = bool(_SKIP_FIG_RE.match(last_nonempty_text)) if last_nonempty_text else False
            if is_skip:
                raw = last_nonempty_text
                skip_label, skip_caption = (raw.split(":", 1) + [""])[:2]
                skip_label = skip_label.strip(); skip_caption = skip_caption.strip()

            # Check if the image paragraph itself carries caption/skip text
            if text:
                if _SKIP_FIG_RE.match(text):
                    is_skip = True
                    skip_label, skip_caption = (text.split(":", 1) + [""])[:2]
                    skip_label = skip_label.strip(); skip_caption = skip_caption.strip()
                elif _FIG_CAPTION_RE.match(text):
                    caption = _strip_fig_label(text)

            if not caption and not is_skip:
                # Look ahead for a caption or skip label (skip up to 2 blanks)
                lookahead = 0
                for _la in range(3):
                    if idx + lookahead < body_len and body_children[idx + lookahead].tag == _W_P:
                        next_text = _para_el_text(body_children[idx + lookahead])
                        if next_text:
                            if _SKIP_FIG_RE.match(next_text):
                                is_skip = True
                                skip_label, skip_caption = (next_text.split(":", 1) + [""])[:2]
                                skip_label = skip_label.strip(); skip_caption = skip_caption.strip()
                                idx_delta = lookahead + 1
                                para_delta = lookahead + 1
                            elif _FIG_CAPTION_RE.match(next_text):
                                caption = _strip_fig_label(next_text)
                                idx_delta = lookahead + 1
                                para_delta = lookahead + 1
                            break
                        lookahead += 1

            # Images in pre_title / authors phases are decorative (cover photos,
            # author headshots) — drop them unless they are skip images.
            # Abstract-phase images (e.g. graphical abstracts after Keywords) ARE
            # included so they appear in the document body.
            if not is_skip and phase in ("pre_title", "authors"):
                if text:
                    last_nonempty_text = text
                continue

            idx      += idx_delta
            para_idx += para_delta

            if is_skip:
                # Include image without Fig-N numbering
                skip_block = _build_figure_block(p, doc, 0, skip_caption)
                skip_block.update({"id": "", "label": skip_label or "Image", "href": ""})
                if current_section is None:
                    current_section = _new_section("", "Other")
                target = (current_section["subsections"][-1]["content"]
                          if current_section["subsections"] else current_section["content"])
                target.append(skip_block)
                continue

            fig_counter += 1
            # Fall back to the pre-scanned caption map if still no caption found
            if not caption and fig_counter in fig_captions:
                caption = _strip_fig_label(fig_captions[fig_counter])
            fig_block = _build_figure_block(p, doc, fig_counter, caption)
            state["figures"].append({k: fig_block[k] for k in ("id", "label", "caption", "href")})
            if current_section is None:
                current_section = _new_section("", "Other")
            target = (current_section["subsections"][-1]["content"]
                      if current_section["subsections"] else current_section["content"])
            target.append(fig_block)
            continue

        if not text:
            continue

        heading_level = _classify_heading(text, style_name, is_bold=is_bold)
        # Strict (no bold heuristic) — used for pre-body phase transitions so that
        # bold decorative lines like "Graphical Abstract" don't skip the real abstract.
        heading_explicit = _classify_heading(text, style_name)
        is_abstract  = bool(_ABSTRACT_RE.match(text))
        is_keywords  = bool(_KEYWORDS_RE.match(text))
        is_refs      = bool(_REFS_HEADING_RE.match(text)) and (heading_level is not None or is_bold)

        # ── pre_title ────────────────────────────────────────────────────────
        if phase == "pre_title":
            if para_idx <= 3 and text.strip().lower().rstrip(':') in _ARTICLE_TYPE_LABELS:
                # "Research paper" / "Review Article" / etc. on its own line
                # above the real title — skip it, stay in pre_title, keep
                # looking for the actual title in the following paragraphs.
                continue
            if heading_explicit == "h1" or (is_bold and font_size and font_size >= 14):
                state["title"] = _para_text_with_fmt(p)
                phase = "authors"
                continue
            if para_idx <= 5:
                state["title"] = _para_text_with_fmt(p)
                phase = "authors"
                continue

        # ── authors ──────────────────────────────────────────────────────────
        if phase == "authors":
            if is_abstract:
                state["authors_raw"] = "\n".join(authors_raw_lines)
                phase = "abstract"
                continue
            # Use strict detection only — bold decorative headings (e.g. "Graphical
            # Abstract") must not prematurely end author collection before the
            # real Abstract heading is encountered.
            #
            # A numbered-heading-shaped line ("N " + capital letter) can also
            # match a numbered affiliation line by coincidence — e.g. "2
            # Department of Physics..." with a stray space after the "2" —
            # which would otherwise end author collection right there and
            # silently drop every affiliation after it (and misroute the
            # real Abstract/body into whatever "section" that line started).
            # Once at least one author line is already collected, a line
            # that itself looks like "<number><affiliation text>" is treated
            # as an affiliation, not a heading, regardless of what
            # _classify_heading thinks.
            looks_like_affil = bool(authors_raw_lines) and bool(_AFFIL_LINE_RE.match(text))
            if heading_explicit and not looks_like_affil:
                state["authors_raw"] = "\n".join(authors_raw_lines)
                phase = "body"
                # intentional fall-through to body processing below
            else:
                authors_raw_lines.append(text)
                continue

        # ── abstract ─────────────────────────────────────────────────────────
        if phase == "abstract":
            if is_keywords:
                kw_part = re.split(r"[:\-]\s*", text, maxsplit=1)[-1]
                state["keywords"] = [k.strip() for k in re.split(r"[,;]", kw_part) if k.strip()]
                phase = "body"
                continue
            if heading_explicit and not is_abstract:
                phase = "body"
                # intentional fall-through — this heading starts the body
                # (use heading_explicit here, not heading_level, so that bold
                # emphasis phrases inside the abstract body don't prematurely
                # end the abstract collection)
            else:
                fmt = _para_text_with_fmt(p)
                state["abstract"] = (state["abstract"] + " " + fmt).strip()
                continue

        # ── body — keyword line after abstract ───────────────────────────────
        if phase == "body" and is_keywords and not state["keywords"]:
            kw_part = re.split(r"[:\-]\s*", text, maxsplit=1)[-1]
            state["keywords"] = [k.strip() for k in re.split(r"[,;]", kw_part) if k.strip()]
            continue

        # ── body → refs transition ───────────────────────────────────────────
        if phase == "body" and is_refs:
            if current_section:
                state["sections"].append(current_section)
                current_section = None
            phase = "refs"
            continue

        # ── refs ─────────────────────────────────────────────────────────────
        if phase == "refs":
            is_list_item = is_list
            is_bullet    = bool(_REF_BULLET_RE.match(text))
            is_numbered  = bool(_REF_ENTRY_RE.match(text))
            starts_lower = bool(text) and text[0].islower()

            fmt = _para_text_with_fmt(p)
            if is_list_item or is_bullet or is_numbered or not state["references"]:
                state["references"].append(fmt)
            elif starts_lower and state["references"]:
                state["references"][-1] += " " + fmt
            else:
                state["references"].append(fmt)
            continue

        # ── body sections ────────────────────────────────────────────────────
        if phase == "body":
            if heading_level in ("h1", "h2"):
                if current_section:
                    state["sections"].append(current_section)
                current_section = _new_section(_para_text_with_fmt(p), _guess_section_type(text))
            elif heading_level == "h3":
                if current_section is None:
                    current_section = _new_section("", "Other")
                current_section["subsections"].append({"heading": _para_text_with_fmt(p), "content": []})
            else:
                if current_section is None:
                    current_section = _new_section("", "Other")

                # Extract paragraph text and check for LaTeX formulas
                para_text = _para_text_with_fmt(p)
                latex_blocks = _extract_latex_blocks(para_text)

                # Add all extracted blocks (text + equations) to content
                target = (current_section["subsections"][-1]["content"]
                         if current_section["subsections"] else current_section["content"])

                for latex_block in latex_blocks:
                    if latex_block["type"] == "text":
                        target.append({"type": "paragraph", "text": latex_block["text"]})
                    else:  # equation
                        target.append({
                            "type": "equation",
                            "text": latex_block["latex"],
                            "latex": latex_block["latex"],
                            "mathml": latex_block["mathml"],
                            "data_uri": ""
                        })

    if current_section:
        state["sections"].append(current_section)


def _new_section(heading: str, sec_type: str) -> dict:
    return {"heading": heading, "type": sec_type, "content": [], "subsections": []}


def _table_to_block(tbl, tbl_number: int) -> dict | None:
    """Convert a python-docx Table object to a content block dict."""
    if not tbl.rows:
        return None
    headers = _row_cells(tbl.rows[0])
    data_rows = [
        _row_cells(r) for r in tbl.rows[1:]
        if any(c.strip() for c in _row_cells(r))
    ]
    label = f"Table {tbl_number}"
    return {
        "type": "table",
        "label": label,
        "caption": label,
        "headers": headers,
        "rows": data_rows,
    }


def _para_el_text(el) -> str:
    """Concatenate all w:t text nodes inside a raw paragraph XML element."""
    return "".join(t.text or "" for t in el.findall(".//" + _W_T)).strip()


def _has_image(p_element) -> bool:
    return bool(
        p_element.findall(f".//{_DRAWING_NS}inline") or
        p_element.findall(f".//{_DRAWING_NS}anchor") or
        p_element.findall(f".//{_VML_NS}imagedata")  # older VML-based images
    )


def _has_math(p_element) -> bool:
    return bool(p_element.findall(f".//{_MATH_NS}oMath"))


def _split_para_math(p):
    """
    Walk a paragraph's direct children in document order, separating plain
    prose text from embedded OMML equations (m:oMath, optionally wrapped in
    m:oMathPara for standalone display equations).

    A paragraph containing math is not always a pure display equation — inline
    equations are commonly used mid-sentence for a single variable (e.g. "...
    where C0- Initial concentration ... and Ct- Concentration..."). Grabbing
    only the first m:oMath and discarding everything else (the old behaviour)
    silently drops the surrounding sentence. This walks every direct child so
    nothing is lost, regardless of how many equations or how much prose the
    paragraph mixes.

    Returns (combined_text, omml_list, has_prose):
      combined_text — full paragraph text with each equation's content spliced
                       in place, using <sub>/<sup> tags for scripts (the same
                       convention _run_text_with_fmt uses for ordinary text)
      omml_list     — OMML XML string for every m:oMath found, in order
      has_prose     — True if any non-whitespace text exists outside the
                       equations, i.e. this is prose with inline math rather
                       than a standalone display equation
    """
    parts = []
    omml_list = []
    has_prose = False

    def _handle_omath(om_el):
        omml_str = _etree.tostring(om_el, encoding='unicode').strip()
        omml_list.append(omml_str)
        parts.append(extract_equation_text(omml_str, html_sub_sup=True))

    for child in p._element:
        tag = child.tag
        if tag == _W_R:
            run_text = _run_text_with_fmt(Run(child, p))
            if run_text.strip():
                has_prose = True
            parts.append(run_text)
        elif tag == _W_HYPERLINK:
            hlink_text = "".join(_run_text_with_fmt(Run(r, p)) for r in child.findall(_W_R))
            if hlink_text.strip():
                has_prose = True
            parts.append(hlink_text)
        elif tag == f"{_MATH_NS}oMath":
            _handle_omath(child)
        elif tag == f"{_MATH_NS}oMathPara":
            for om in child.findall(f"{_MATH_NS}oMath"):
                _handle_omath(om)
        # else: paragraph properties, proofErr markers, bookmarks, etc. — no text

    return "".join(parts).strip(), omml_list, has_prose


def _math_para_to_image(p) -> str:
    """Render an OMML equation paragraph to a base64 PNG via LibreOffice."""
    import subprocess, tempfile, os, glob, shutil, base64 as _b64, io
    from docx import Document as _Doc
    from copy import deepcopy
    from PIL import Image

    tmp_doc = _Doc()
    # Remove the default blank paragraph so only the equation appears
    for existing in tmp_doc.paragraphs:
        existing._element.getparent().remove(existing._element)
    tmp_doc.element.body.insert(0, deepcopy(p._element))

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        tmp_doc.save(f.name)
        src = f.name

    out_dir = tempfile.mkdtemp()
    try:
        subprocess.run(
            ["soffice", "--headless", "--convert-to", "png", src, "--outdir", out_dir],
            check=True, capture_output=True, timeout=30
        )
        pngs = glob.glob(os.path.join(out_dir, "*.png"))
        if pngs:
            img = Image.open(pngs[0])
            if img.mode in ("RGBA", "LA") or (img.mode == "P" and "transparency" in img.info):
                # Same transparency-compositing fix as _blob_to_data_uri(): a
                # plain .convert("RGB") drops the alpha channel without
                # compositing, exposing whatever raw color sits under the
                # "transparent" area instead of rendering it as white.
                img = img.convert("RGBA")
                background = Image.new("RGB", img.size, (255, 255, 255))
                background.paste(img, mask=img.split()[-1])
                img = background
            else:
                img = img.convert("RGB")
            img = _smart_crop(img)
            buf = io.BytesIO()
            img.save(buf, format="PNG")
            return f"data:image/png;base64,{_b64.b64encode(buf.getvalue()).decode()}"
    except Exception:
        pass
    finally:
        os.unlink(src)
        shutil.rmtree(out_dir, ignore_errors=True)
    return ""


def _smart_crop(img, pad: int = 8, min_blank_run: int = 30):
    """
    Crop to actual content area, skipping thin border-only rows/columns.

    Word EMF drawing canvases often have a 1–3 px border around a large blank
    area with the chart content in one corner.  A simple bounding-box crop
    keeps the border (and the blank interior it encloses) because the border
    pixels extend to all four edges of the canvas.

    This function finds rows/columns whose non-white pixel count exceeds a
    threshold sized relative to the image dimensions.  Only crops a given side
    if the blank run on that side is at least min_blank_run pixels — this
    prevents over-cropping images where the content already fills most of the
    canvas (e.g. charts with thick outer border frames near the image edge).
    """
    try:
        import numpy as np
        arr = np.array(img.convert("RGB"))
        h, w = arr.shape[:2]
        non_white = (arr < 250).any(axis=2)          # True where pixel ≠ white
        row_density = non_white.sum(axis=1).astype(int)
        col_density = non_white.sum(axis=0).astype(int)

        # Minimum pixels per row/col to be counted as real content.
        # A 1-3 px border contributes ≤ 6 px; chart content contributes much more.
        min_px = max(6, min(w, h) // 30)

        cr = np.where(row_density >= min_px)[0]
        cc = np.where(col_density >= min_px)[0]
        if not len(cr) or not len(cc):
            # Fallback: any non-white pixel
            cr = np.where(row_density > 0)[0]
            cc = np.where(col_density > 0)[0]
        if not len(cr) or not len(cc):
            return img

        # Only crop a side if the blank run there is large enough to be real
        # canvas waste (not just a 1-2 px anti-alias fringe or rounding gap).
        top    = max(0, int(cr[0])  - pad) if cr[0]       > min_blank_run else 0
        bottom = min(h, int(cr[-1]) + pad + 1) if (h - 1 - cr[-1]) > min_blank_run else h
        left   = max(0, int(cc[0])  - pad) if cc[0]       > min_blank_run else 0
        right  = min(w, int(cc[-1]) + pad + 1) if (w - 1 - cc[-1]) > min_blank_run else w

        return img.crop((left, top, right, bottom))
    except ImportError:
        # numpy unavailable — fall back to simple bounding-box crop
        from PIL import ImageChops
        bg = img._new(img.mode, img.size)
        bg.paste((255, 255, 255), (0, 0, img.width, img.height))
        bbox = ImageChops.difference(img.convert("RGB"),
                                     img.convert("RGB").point(lambda _: 255)).getbbox()
        if not bbox:
            return img
        l, t, r, b = bbox
        return img.crop((max(0, l - pad), max(0, t - pad),
                         min(img.width, r + pad), min(img.height, b + pad)))


_WEASYPRINT_OK_TYPES = {"image/png", "image/jpeg", "image/jpg", "image/gif",
                        "image/svg+xml", "image/webp", "image/bmp"}


def _blob_to_pil_image(blob: bytes, content_type: str):
    """Decode an image blob to a PIL Image, converting WMF/EMF to a raster
    via LibreOffice (falling back to Pillow directly for anything else, or
    if LibreOffice isn't available). Returns None on failure.

    Used both by _blob_to_data_uri() (single-image fast path) and by
    _blobs_to_data_uri() (which needs real Image objects to composite
    multiple sub-panel images side-by-side)."""
    from PIL import Image
    import io
    ct = (content_type or "").lower()
    if ct in _WEASYPRINT_OK_TYPES:
        try:
            return Image.open(io.BytesIO(blob)).convert("RGB")
        except Exception:
            return None
    # Use LibreOffice headless for EMF/WMF (most reliable cross-platform renderer)
    try:
        import subprocess, tempfile, os, glob, shutil
        suffix = ".emf" if "emf" in ct else ".wmf"
        with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as src:
            src.write(blob)
            src_path = src.name
        out_dir = tempfile.mkdtemp()
        try:
            subprocess.run(
                ["soffice", "--headless", "--convert-to", "png", src_path, "--outdir", out_dir],
                check=True, capture_output=True, timeout=30
            )
            pngs = glob.glob(os.path.join(out_dir, "*.png"))
            if pngs:
                img = Image.open(pngs[0])
                if img.mode in ("RGBA", "LA") or (img.mode == "P" and "transparency" in img.info):
                    # LibreOffice's EMF/WMF rendering often leaves the area
                    # outside the actual chart transparent. Plain
                    # .convert("RGB") on an RGBA image just drops the alpha
                    # channel — it does NOT composite onto white — so
                    # whatever raw color happens to sit under those
                    # "invisible" pixels (frequently a solid blue/black
                    # placeholder from the renderer) becomes visible as a
                    # wrong-colored background fill. Composite onto white
                    # first so transparent areas render as white, like the
                    # original chart actually looks.
                    img = img.convert("RGBA")
                    background = Image.new("RGB", img.size, (255, 255, 255))
                    background.paste(img, mask=img.split()[-1])
                    img = background
                else:
                    img = img.convert("RGB")
                img.load()  # fully read pixel data before the temp dir is removed
                return _smart_crop(img)
        finally:
            os.unlink(src_path)
            shutil.rmtree(out_dir, ignore_errors=True)
    except Exception:
        pass
    # Pillow fallback (works for some formats on Windows)
    try:
        return Image.open(io.BytesIO(blob)).convert("RGB")
    except Exception:
        return None


def _blob_to_data_uri(blob: bytes, content_type: str) -> str:
    """Return a base64 data URI, converting WMF/EMF to PNG via LibreOffice or Pillow."""
    import base64 as _b64
    ct = (content_type or "").lower()
    if ct in _WEASYPRINT_OK_TYPES:
        # Fast path: pass the original bytes straight through so an
        # already web-safe image (PNG/JPEG) isn't needlessly re-encoded.
        b64 = _b64.b64encode(blob).decode()
        return f"data:{content_type};base64,{b64}"
    img = _blob_to_pil_image(blob, content_type)
    if img is None:
        return ""
    import io
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    b64 = _b64.b64encode(buf.getvalue()).decode()
    return f"data:image/png;base64,{b64}"


def _hconcat_images(images: list, gap: int = 16):
    """Concatenate images left-to-right on a white background, vertically
    centered on the tallest one. Used for figures whose sub-panels — e.g.
    (a)/(b) — are embedded in the DOCX as separate side-by-side images
    sharing one caption, rather than as a single combined picture."""
    from PIL import Image
    total_w = sum(im.width for im in images) + gap * (len(images) - 1)
    max_h = max(im.height for im in images)
    canvas = Image.new("RGB", (total_w, max_h), (255, 255, 255))
    x = 0
    for im in images:
        y = (max_h - im.height) // 2
        canvas.paste(im, (x, y))
        x += im.width + gap
    return canvas


def _blobs_to_data_uri(blobs: list) -> str:
    """Encode one or more (blob, content_type) image tuples as a single
    data URI. A lone image takes the normal single-image path. Multiple
    images (a figure with separately-embedded sub-panels — see
    _hconcat_images) are decoded and composited side-by-side into one
    image, so a figure isn't silently truncated to just its first panel."""
    if not blobs:
        return ""
    if len(blobs) == 1:
        return _blob_to_data_uri(*blobs[0])

    images = [img for img in (_blob_to_pil_image(b, ct) for b, ct in blobs) if img is not None]
    if not images:
        return ""

    import base64 as _b64, io
    composite = images[0] if len(images) == 1 else _hconcat_images(images)
    buf = io.BytesIO()
    composite.save(buf, format="PNG")
    return f"data:image/png;base64,{_b64.b64encode(buf.getvalue()).decode()}"


def _read_xfrm_chOffExt(grp_elem):
    """Read a wpg:wgp/wpg:grpSp element's own a:chOff/a:chExt — the
    logical coordinate space its direct pic:pic children's <a:off>/<a:ext>
    are expressed in. Returns (chOffX, chOffY, chExtCx, chExtCy) or None."""
    grpSpPr = grp_elem.find(f"{_WPG_NS}grpSpPr")
    if grpSpPr is None:
        return None
    xfrm = grpSpPr.find(f"{_DRAW_A_NS}xfrm")
    if xfrm is None:
        return None
    chOff = xfrm.find(f"{_DRAW_A_NS}chOff")
    chExt = xfrm.find(f"{_DRAW_A_NS}chExt")
    if chOff is None or chExt is None:
        return None
    try:
        return (int(chOff.get("x")), int(chOff.get("y")),
                int(chExt.get("cx")), int(chExt.get("cy")))
    except (TypeError, ValueError):
        return None


def _walk_group_pics(grp_elem, abs_x: float, abs_y: float, abs_cx: float, abs_cy: float) -> list:
    """Recursively collect every pic:pic inside grp_elem — at ANY nesting
    depth, flattened — with its position/size mapped into one absolute
    coordinate space. Word groups can nest (e.g. two charts grouped
    together as one sub-group, with a data-table picture placed as a
    direct sibling of that sub-group rather than inside it — exactly the
    shape of a "chart with an overlaid results table" composite figure),
    and each nesting level has its own internal a:chOff/a:chExt coordinate
    system, so a picture's raw <a:off>/<a:ext> is meaningless without
    composing every ancestor group's transform down to one common frame.

    (abs_x, abs_y, abs_cx, abs_cy) is where grp_elem's *own* chOff/chExt
    origin+extent maps to in that common frame — pass grp_elem's own
    chOff/chExt right back for the top-level call (identity transform).

    Returns [(r_embed, x, y, cx, cy), ...] in the common frame's units.
    """
    grpSpPr = grp_elem.find(f"{_WPG_NS}grpSpPr")
    xfrm = grpSpPr.find(f"{_DRAW_A_NS}xfrm") if grpSpPr is not None else None
    if xfrm is None:
        return []
    chOff = xfrm.find(f"{_DRAW_A_NS}chOff")
    chExt = xfrm.find(f"{_DRAW_A_NS}chExt")
    if chOff is None or chExt is None:
        return []
    try:
        chOffX, chOffY = int(chOff.get("x")), int(chOff.get("y"))
        chExtCx, chExtCy = int(chExt.get("cx")), int(chExt.get("cy"))
    except (TypeError, ValueError):
        return []
    if chExtCx <= 0 or chExtCy <= 0:
        return []
    scale_x = abs_cx / chExtCx
    scale_y = abs_cy / chExtCy

    def to_abs(x, y, cx, cy):
        return (abs_x + (x - chOffX) * scale_x, abs_y + (y - chOffY) * scale_y,
                cx * scale_x, cy * scale_y)

    # Walk grp_elem's direct children IN DOCUMENT ORDER (paint order) —
    # not pic:pic-then-grpSp in two separate passes, which would silently
    # reorder them whenever a nested group appears *before* a sibling
    # picture in the source (as it does here: two charts grouped together
    # come first, then a data-table picture on top of each — processing
    # all direct pics before recursing into the group would paint the
    # charts last, burying the tables underneath them).
    results = []
    pic_tag = f"{_PIC_NS}pic"
    grpsp_tag = f"{_WPG_NS}grpSp"
    for child in grp_elem:
        if child.tag == pic_tag:
            blip = child.find(f".//{_DRAW_A_NS}blip")
            if blip is None:
                continue
            r_embed = blip.get(f"{_REL_NS}embed")
            if not r_embed:
                continue
            spPr = child.find(f"{_PIC_NS}spPr")
            pxfrm = spPr.find(f"{_DRAW_A_NS}xfrm") if spPr is not None else None
            poff = pxfrm.find(f"{_DRAW_A_NS}off") if pxfrm is not None else None
            pext = pxfrm.find(f"{_DRAW_A_NS}ext") if pxfrm is not None else None
            if poff is None or pext is None:
                continue
            try:
                x, y = int(poff.get("x")), int(poff.get("y"))
                cx, cy = int(pext.get("cx")), int(pext.get("cy"))
            except (TypeError, ValueError):
                continue
            results.append((r_embed, *to_abs(x, y, cx, cy)))

        elif child.tag == grpsp_tag:
            sub_grpSpPr = child.find(f"{_WPG_NS}grpSpPr")
            sub_xfrm = sub_grpSpPr.find(f"{_DRAW_A_NS}xfrm") if sub_grpSpPr is not None else None
            sub_off = sub_xfrm.find(f"{_DRAW_A_NS}off") if sub_xfrm is not None else None
            sub_ext = sub_xfrm.find(f"{_DRAW_A_NS}ext") if sub_xfrm is not None else None
            if sub_off is None or sub_ext is None:
                continue
            try:
                sx, sy = int(sub_off.get("x")), int(sub_off.get("y"))
                scx, scy = int(sub_ext.get("cx")), int(sub_ext.get("cy"))
            except (TypeError, ValueError):
                continue
            if scx <= 0 or scy <= 0:
                continue
            asx, asy, ascx, ascy = to_abs(sx, sy, scx, scy)
            if ascx <= 0 or ascy <= 0:
                continue
            results.extend(_walk_group_pics(child, asx, asy, ascx, ascy))

    return results


def _extract_grouped_pic_specs(drawing) -> list:
    """If `drawing` (a wp:inline/wp:anchor element) contains a Word GROUP
    shape (wpg:wgp — multiple pictures the author selected and grouped
    into one object, e.g. a data table image overlaid on a chart image to
    build one composite figure panel, possibly with further nested
    sub-groups), return every picture found at any nesting depth with its
    position/size as fractions of the outermost group's coordinate space:
    [{"r_embed", "fx", "fy", "fw", "fh"}, ...], in document order (=
    paint/z order — later entries are drawn on top). Returns [] if
    `drawing` isn't a usable group, so the caller falls back to treating
    each picture as an independent, unpositioned image."""
    for grp in drawing.findall(f".//{_WPG_NS}wgp"):
        coord = _read_xfrm_chOffExt(grp)
        if not coord:
            continue
        chOffX, chOffY, chExtCx, chExtCy = coord
        if chExtCx <= 0 or chExtCy <= 0:
            continue
        flat = _walk_group_pics(grp, chOffX, chOffY, chExtCx, chExtCy)
        if len(flat) < 2:
            continue
        return [
            {
                "r_embed": r_embed,
                "fx": (x - chOffX) / chExtCx,
                "fy": (y - chOffY) / chExtCy,
                "fw": cx / chExtCx,
                "fh": cy / chExtCy,
            }
            for r_embed, x, y, cx, cy in flat
        ]
    return []


def _positioned_composite_to_data_uri(items: list, canvas_width: int = 1600) -> str:
    """Composite pictures onto one canvas using each one's fractional
    position/size within the group's overall bounding box, painting in
    document order so a later item (e.g. a data table overlaid on a
    chart) lands on top of an earlier one instead of beside it."""
    from PIL import Image
    import base64 as _b64, io

    max_x = max((it["fx"] + it["fw"] for it in items), default=0)
    max_y = max((it["fy"] + it["fh"] for it in items), default=0)
    if max_x <= 0 or max_y <= 0:
        return ""
    canvas_height = max(1, int(canvas_width * max_y / max_x))
    canvas = Image.new("RGB", (canvas_width, canvas_height), (255, 255, 255))

    pasted = 0
    for it in items:
        img = _blob_to_pil_image(it["blob"], it["content_type"])
        if img is None:
            continue
        w = max(1, round(it["fw"] / max_x * canvas_width))
        h = max(1, round(it["fh"] / max_y * canvas_height))
        x = round(it["fx"] / max_x * canvas_width)
        y = round(it["fy"] / max_y * canvas_height)
        img = img.resize((w, h))
        if img.mode in ("RGBA", "LA") or (img.mode == "P" and "transparency" in img.info):
            img = img.convert("RGBA")
            canvas.paste(img, (x, y), mask=img.split()[-1])
        else:
            canvas.paste(img.convert("RGB"), (x, y))
        pasted += 1

    if pasted == 0:
        return ""
    buf = io.BytesIO()
    canvas.save(buf, format="PNG")
    return f"data:image/png;base64,{_b64.b64encode(buf.getvalue()).decode()}"


def _build_figure_block(p, doc, fig_number: int, caption: str) -> dict:
    """Extract image bytes (base64 data URI) and return a figure content
    block. A figure paragraph can embed more than one image — e.g.
    side-by-side (a)/(b) sub-panels pasted as two separate pictures under
    one shared caption — so every image found is collected and, when
    there's more than one, composited side-by-side (_blobs_to_data_uri)
    rather than keeping only the first and silently dropping the rest."""
    blobs = []  # [(blob, content_type), ...] in document order

    # 1) DrawingML (modern DOCX: wp:inline / wp:anchor). A single
    # wp:inline/wp:anchor can itself be a *group* shape (wpg:wgp, e.g.
    # Word's "Group" command used to combine sub-panels — or a chart
    # picture with a data-table picture overlaid on top of it — into one
    # object) containing several pic:pic/blip elements. Try the
    # position-aware group path first (it knows where each picture goes,
    # e.g. "table overlaid in the chart's corner" vs. "beside it"); if a
    # drawing isn't a group, fall back to collecting every blip in it
    # (plain .find() would silently keep only the first and drop the rest).
    group_items = None
    for tag in (f"{_DRAWING_NS}inline", f"{_DRAWING_NS}anchor"):
        for drawing in p._element.findall(f".//{tag}"):
            specs = _extract_grouped_pic_specs(drawing)
            if specs:
                group_items = specs
                continue
            for blip in drawing.findall(f".//{_DRAW_A_NS}blip"):
                r_embed = blip.get(f"{_REL_NS}embed")
                if not r_embed:
                    continue
                try:
                    img_part = doc.part.related_parts[r_embed]
                    blobs.append((img_part.blob, img_part.content_type))
                except Exception:
                    pass

    # 2) VML fallback (older DOCX: v:imagedata r:id="rId..."), only tried
    # when DrawingML found nothing — same precedence as before.
    if not blobs and not group_items:
        for imgdata in p._element.findall(f".//{_VML_NS}imagedata"):
            r_id = imgdata.get(f"{_REL_NS}id") or imgdata.get(f"{_O_NS}relid")
            if not r_id:
                continue
            try:
                img_part = doc.part.related_parts[r_id]
                blobs.append((img_part.blob, img_part.content_type))
            except Exception:
                pass

    if group_items:
        for item in group_items:
            try:
                img_part = doc.part.related_parts[item["r_embed"]]
                item["blob"] = img_part.blob
                item["content_type"] = img_part.content_type
            except Exception:
                item["blob"] = None
        group_items = [it for it in group_items if it["blob"] is not None]
        data_uri = _positioned_composite_to_data_uri(group_items) if group_items else ""
        if not data_uri:
            data_uri = _blobs_to_data_uri(blobs)  # last-resort fallback
    else:
        data_uri = _blobs_to_data_uri(blobs)

    return {
        "type":     "figure",
        "id":       f"fig{fig_number}",
        "label":    f"Figure {fig_number}",
        "caption":  caption,
        "href":     f"figure{fig_number}.png",
        "data_uri": data_uri,
    }


def _guess_section_type(heading: str) -> str:
    """Map heading text to one of the JATS section type labels."""
    h = heading.lower()
    mapping = [
        ("introduction",   "Introduction"),
        ("background",     "Introduction"),
        ("literature",     "Introduction"),
        ("material",       "Methods"),
        ("method",         "Methods"),
        ("experiment",     "Methods"),
        ("result",         "Results"),
        ("finding",        "Results"),
        ("discussion",     "Discussion"),
        ("conclusion",     "Conclusion"),
        ("summary",        "Conclusion"),
        ("acknowledg",     "Acknowledgements"),
        ("funding",        "Acknowledgements"),
    ]
    for key, val in mapping:
        if key in h:
            return val
    return "Other"


def _parse_references(raw_refs: list) -> list:
    """
    Convert raw collected reference strings into structured dicts.

    Handles three input formats (stripped / cleaned here):
      • Bullet/dash list:  "- Ibrahim Khan, Khalid Saeed... doi: 10.1016/..."
      • Numbered [1]:      "[1] Smith J et al. J Chem. 2020..."
      • Plain paragraph:   "Patel R. Green chemistry. Org Lett. 2021..."

    Extracts a DOI from text when present. References without one keep doi=""
    — use POST /enrich-refs (Crossref lookup) to fill those in afterward;
    that lookup is intentionally not done here since it's a slow, per-reference
    external API call that would block the whole /parse request.
    Returns: [{number, raw_text, doi}]
    """
    result = []
    for i, raw in enumerate(raw_refs, 1):
        # Strip leading bullet character or number marker
        text = _REF_BULLET_RE.sub("", raw)
        text = re.sub(r"^\[?\d+[\]\.]\s+", "", text).strip()

        # Extract DOI from text (supports: "doi: 10.xxx", "https://doi.org/10.xxx", bare "10.xxxx/yyyy")
        doi = ""
        m = _DOI_RE.search(raw)
        if m:
            doi_str = m.group(1).rstrip(".,;)>").strip()
            # Remove trailing punctuation and whitespace that's common in text
            doi_str = re.sub(r'[\s\.\-_]+$', '', doi_str).strip()
            if doi_str:
                doi = doi_str

        if text:
            result.append({"number": i, "raw_text": text, "doi": doi})

    return result


def _parse_authors(raw: str) -> list:
    """
    Parse the author block captured between the title and abstract.

    Handles the NFP manuscript format:
      Line 0:  Mylarappa M¹*, N Raghavendra²*, Shravan Kumar K N³, Chandruavasan S⁴
      Line 1+: ¹Department of Studies in Chemistry, Bangalore University…
               ²Research Centre, Department of Chemistry…
               Corresponding authors: email1@x.com, email2@x.com
    """
    if not raw:
        return [_empty_author()]

    lines = [l.strip() for l in raw.split("\n") if l.strip()]
    if not lines:
        return [_empty_author()]

    affil_map = _build_affil_map(lines[1:])
    corresp_emails = _find_corresp_emails(lines)

    authors = _split_author_line(lines[0])
    if not authors:
        return [_empty_author()]

    # Assign affiliations and emails
    email_idx = 0
    for author in authors:
        sup_nums = author.pop("_sup_nums", [])
        # An author affiliated with more than one institution (e.g. "5,6")
        # gets each affiliation's text, deduplicated, as a real list —
        # "affiliations" is the source of truth (each entry gets its own
        # numbered superscript in exports); "affiliation" is kept as a
        # joined-string convenience/legacy field for any code not yet
        # updated to read the list.
        affs = []
        for n in sup_nums:
            aff = affil_map.get(n, "")
            if aff and aff not in affs:
                affs.append(aff)
        author["affiliations"] = affs
        author["affiliation"] = "; ".join(affs)
        if author["corresponding"]:
            author["email"] = corresp_emails[email_idx] if email_idx < len(corresp_emails) else ""
            email_idx += 1

    return authors


def _empty_author() -> dict:
    return {
        "first_name": "", "last_name": "", "affiliation": "", "affiliations": [],
        "email": "", "orcid": "", "corresponding": True,
    }


def _split_author_line(line: str) -> list:
    """
    Split an author name line on superscript markers (Unicode ¹²³⁴ or
    digits glued directly to a letter like M1*).

    Returns a list of author dicts with a temporary '_sup_nums' key holding
    the list of normalised affiliation-number strings (usually one, but an
    author affiliated with more than one institution has several, e.g.
    "Chan5,6" -> ["5", "6"]).
    """
    print(f"DEBUG _split_author_line: input = {line[:100]}...")
    parts = _AUTHOR_MARKER_RE.split(line)
    print(f"DEBUG _split_author_line: split into {len(parts)} parts")
    for idx, p in enumerate(parts[:10]):
        print(f"  parts[{idx}] = {p[:60]}...")

    # No markers found — fall back to simple comma split
    if len(parts) == 1:
        print("DEBUG: No markers found, using fallback comma split")
        return _fallback_comma_split(line)

    authors = []
    for i in range(0, len(parts), 2):
        name_raw = parts[i].strip().lstrip(',').strip()
        marker = parts[i + 1] if i + 1 < len(parts) else ""

        if not name_raw:
            continue

        is_corresp = '*' in marker
        # Strip the asterisk and any trailing/leading separator comma, then
        # split the remainder on ',' so "5,6" becomes ["5", "6"] instead of
        # collapsing into the single (nonexistent) affiliation number "56".
        marker_nums = marker.replace('*', '').strip(',').translate(_SUP_TO_NUM)
        sup_nums = [n for n in marker_nums.split(',') if n]

        first, last = _split_name(name_raw)
        print(f"DEBUG: Author {i//2 + 1}: first='{first}' last='{last}' sup={sup_nums} corresp={is_corresp}")
        authors.append({
            "first_name": first,
            "last_name": last,
            "affiliation": "",
            "email": "",
            "orcid": "",
            "corresponding": is_corresp,
            "_sup_nums": sup_nums,
        })

    print(f"DEBUG: _split_author_line returning {len(authors)} authors")
    return authors or _fallback_comma_split(line)


def _split_name(name: str) -> tuple:
    """
    Return (first_name, last_name).

    Trailing single-letter initials become the last name so that South-Indian /
    abbreviated suffix formats are handled correctly:
      'Shankar S'          → ('Shankar',       'S')
      'Ravi K N'           → ('Ravi',           'K N')
      'Manju Kumar S N'    → ('Manju Kumar',    'S N')
      'Mylarappa M'        → ('Mylarappa',      'M')
      'Shravan Kumar K N'  → ('Shravan Kumar',  'K N')

    If the final token is a full word it is the last name as normal:
      'S.K RaviKumar'      → ('S.K',            'RaviKumar')
      'John Smith'         → ('John',            'Smith')

    A single-word name has no last name:
      'Ravi'               → ('Ravi',            '')
    """
    parts = name.strip().split()
    if not parts:
        return ("", "")
    if len(parts) == 1:
        # Handle "S.N.Manjula" — period-separated initials followed by a full word
        dot_parts = [s for s in parts[0].split(".") if s]
        if len(dot_parts) >= 2:
            *initials, last_part = dot_parts
            if (all(len(i) == 1 and i.isalpha() for i in initials)
                    and len(last_part) > 1 and last_part.isalpha()):
                return (".".join(initials) + ".", last_part)
        return (parts[0], "")

    def _is_initial(token: str) -> bool:
        """True for single-letter tokens like 'K', 'N', 'S', 'M.' (period optional)."""
        return len(token.rstrip(".")) == 1 and token.rstrip(".").isalpha()

    # Walk backwards collecting consecutive single-letter initials into last_name
    tail_start = len(parts)
    while tail_start > 0 and _is_initial(parts[tail_start - 1]):
        tail_start -= 1

    if 0 < tail_start < len(parts):
        # Non-initial first-name part(s) + trailing initials
        return (" ".join(parts[:tail_start]), " ".join(parts[tail_start:]))

    # No trailing initials (or every token is an initial) — last token is last name
    return (" ".join(parts[:-1]), parts[-1])


def _build_affil_map(lines: list) -> dict:
    """Build {number_str: affiliation_text} from lines like '¹Department of…'"""
    affil_map = {}
    for line in lines:
        m = re.match(r'^([¹²³⁴⁵⁶⁷⁸⁹⁰]+|\d+)\s*(.+)$', line)
        if m:
            num = m.group(1).translate(_SUP_TO_NUM)
            affil_map[num] = m.group(2).strip()
    return affil_map


def _find_corresp_emails(lines: list) -> list:
    """Extract emails from a 'Corresponding author(s): …' line if present,
    otherwise collect all emails in the block."""
    for line in lines:
        if _CORRESP_LINE_RE.search(line):
            return _EMAIL_RE.findall(line)
    # Fallback: any email anywhere in the author block
    emails = []
    for line in lines:
        emails.extend(_EMAIL_RE.findall(line))
    return emails


def _fallback_comma_split(line: str) -> list:
    """Simple comma split used when no superscript markers are detected."""
    print(f"DEBUG _fallback_comma_split: input = {line[:100]}...")
    names = [n.strip() for n in line.split(',') if n.strip()]
    print(f"DEBUG _fallback_comma_split: split into {len(names)} names on commas")
    result = []
    for i, name in enumerate(names):
        name_clean = re.sub(r'[¹²³⁴⁵⁶⁷⁸⁹⁰\d*]', '', name).strip()
        first, last = _split_name(name_clean)
        print(f"DEBUG: Fallback Author {i+1}: raw='{name}' clean='{name_clean}' first='{first}' last='{last}'")
        result.append({
            "first_name": first,
            "last_name": last,
            "affiliation": "",
            "email": "",
            "orcid": "",
            "corresponding": i == 0,
            "_sup_nums": [str(i + 1)],
        })
    print(f"DEBUG _fallback_comma_split: returning {len(result)} authors")
    return result


def _row_cells(row) -> list:
    """Return cell texts for a table row, skipping duplicate merged-cell references."""
    seen  = set()
    cells = []
    for cell in row.cells:
        cid = id(cell._tc)
        if cid not in seen:
            seen.add(cid)
            cells.append(_cell_text_with_fmt(cell))
    return cells


