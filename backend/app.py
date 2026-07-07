"""
NFP Article Formatter — Flask REST API
Endpoints: POST /parse, POST /generate, POST /validate, POST /autotag
"""
import os
import tempfile
import json
import anthropic
import logging

import re
from flask import Flask, request, jsonify, make_response
from flask_cors import CORS
from dotenv import load_dotenv

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

from parser.docx_parser import parse_docx
from jats.generator import generate_jats
from jats.validator import validate_jats

load_dotenv()

app = Flask(__name__)
CORS(app, resources={r"/*": {"origins": "*"}})

ALLOWED_EXTENSIONS = {"docx"}
MAX_CONTENT_LENGTH = 200 * 1024 * 1024  # 200 MB (for poster images)
app.config["MAX_CONTENT_LENGTH"] = MAX_CONTENT_LENGTH


def allowed_file(filename: str) -> bool:
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS


@app.route("/health", methods=["GET"])
def health():
    return jsonify({"status": "ok"})


@app.route("/parse", methods=["POST"])
def parse():
    if "file" not in request.files:
        return jsonify({"error": "No file part in request"}), 400

    file = request.files["file"]
    if file.filename == "":
        return jsonify({"error": "No file selected"}), 400

    if not allowed_file(file.filename):
        return jsonify({"error": "Only .docx files are supported"}), 400

    doc_mode = request.form.get("doc_mode", "article")  # article | abstracts | poster

    try:
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            file.save(tmp.name)
            tmp_path = tmp.name

        if doc_mode in ("abstracts", "poster_abstracts"):
            from parser.abstract_parser import parse_abstract_collection
            result = parse_abstract_collection(tmp_path)
            result["collection_type"] = doc_mode  # "abstracts" or "poster_abstracts"
        elif doc_mode == "poster":
            from parser.poster_parser import parse_poster
            result = parse_poster(tmp_path)
        else:
            result = parse_docx(tmp_path)
        return jsonify(result)
    except Exception as e:
        return jsonify({"error": f"Parsing failed: {str(e)}"}), 500
    finally:
        try:
            os.unlink(tmp_path)
        except Exception:
            pass


@app.route("/export/abstracts-xml", methods=["POST"])
def export_abstracts_xml():
    """
    Generate an OJS Native XML file (native.xsd) containing all abstracts,
    ready for import via OJS → Tools → Import/Export → Native XML Plugin.

    Root element: <articles xmlns="http://pkp.sfu.ca" ...>
    Each abstract becomes an <article> with locale-tagged title, abstract,
    keywords, and authors using OJS's givenname/familyname structure.
    """
    from xml.etree.ElementTree import Element, SubElement, tostring
    from xml.dom import minidom

    data = request.get_json(force=True)
    if not data or "abstracts" not in data:
        return jsonify({"error": "Expected JSON body with 'abstracts' key"}), 400

    abstracts  = data["abstracts"]
    event_name = data.get("event_name", data.get("doc_title", "Conference"))
    year       = str(data.get("year", "2025"))
    locale     = data.get("locale", "en")
    section    = data.get("section_ref", "ABS")

    def _txt(el, val):
        el.text = str(val or "")
        return el

    def _build_article(root_el: Element, ab: dict, seq: int):
        # OJS 3.x structure:
        #   <article>          ← submission shell (no content attributes)
        #     <id/>
        #     <publication>    ← all metadata lives here
        #       <id/>
        #       <title/>
        #       <abstract/>
        #       <keywords/>
        #       <authors/>
        #     </publication>
        #   </article>

        article = SubElement(root_el, "article")
        article.set("locale",              locale)
        article.set("submission_progress", "0")
        article.set("stage",               "production")

        _txt(SubElement(article, "id", {"type": "internal", "advice": "ignore"}), str(seq))

        # <publication> - per native.xsd, section_ref is required
        # seq, primary_contact_id, access_status are optional but recommended
        pub = SubElement(article, "publication")
        pub.set("version",            "1")
        pub.set("section_ref",        section)
        pub.set("seq",                str(seq - 1))
        pub.set("primary_contact_id", "0")
        pub.set("access_status",      "0")

        _txt(SubElement(pub, "id", {"type": "internal", "advice": "ignore"}), str(seq))

        # Issue identification - OJS requires this to link publication to an issue
        issue_id = SubElement(pub, "issue_identification")
        _txt(SubElement(issue_id, "volume"), "1")
        _txt(SubElement(issue_id, "number"), "1")
        _txt(SubElement(issue_id, "year"),   year)

        # Title
        title_el = SubElement(pub, "title")
        title_el.set("locale", locale)
        title_el.text = (ab.get("title") or f"Abstract {seq}").strip()

        # Abstract
        abstract_text = (ab.get("abstract") or "").strip()
        if abstract_text:
            ab_el = SubElement(pub, "abstract")
            ab_el.set("locale", locale)
            ab_el.text = abstract_text

        # Keywords
        kws = ab.get("keywords") or []
        if kws:
            kwd_group = SubElement(pub, "keywords")
            kwd_group.set("locale", locale)
            for kw in kws:
                _txt(SubElement(kwd_group, "keyword"), kw.strip())

        # Authors
        authors = ab.get("authors") or []
        if authors:
            authors_el = SubElement(pub, "authors")
            for i, a in enumerate(authors):
                author_el = SubElement(authors_el, "author")
                author_el.set("include_in_browse", "true")
                author_el.set("user_group_ref",    "Author")
                author_el.set("seq",               str(i))
                author_el.set("id",                str(i + 1))

                given_el = SubElement(author_el, "givenname")
                given_el.set("locale", locale)
                given_el.text = (a.get("first_name") or "").strip()

                family_el = SubElement(author_el, "familyname")
                family_el.set("locale", locale)
                family_el.text = (a.get("last_name") or "").strip()

                aff = (a.get("affiliation") or "").strip()
                if aff:
                    aff_el = SubElement(author_el, "affiliation")
                    aff_el.set("locale", locale)
                    aff_el.text = aff

                # <email> is required by native.xsd (one of affiliation/country/email
                # must be present) — always include it even if empty
                _txt(SubElement(author_el, "email"), (a.get("email") or "").strip())

    # ── Build root <articles> with OJS namespace ─────────────────────────────
    root = Element("articles")
    root.set("xmlns",              "http://pkp.sfu.ca")
    root.set("xmlns:xsi",         "http://www.w3.org/2001/XMLSchema-instance")
    root.set("xsi:schemaLocation", "http://pkp.sfu.ca native.xsd")

    for i, ab in enumerate(abstracts, 1):
        try:
            _build_article(root, ab, i)
        except Exception:
            pass

    raw    = tostring(root, encoding="unicode", xml_declaration=False)
    pretty = minidom.parseString(raw).toprettyxml(indent="  ")
    pretty = "\n".join(pretty.split("\n")[1:])   # strip minidom XML declaration
    xml_str = '<?xml version="1.0" encoding="UTF-8"?>\n' + pretty

    slug = re.sub(r"[^a-zA-Z0-9_\-]", "_", event_name)[:60].strip("_") or "abstracts"
    resp = make_response(xml_str.encode("utf-8"))
    resp.headers["Content-Type"] = "application/xml; charset=utf-8"
    resp.headers["Content-Disposition"] = f'attachment; filename="{slug}_ojs.xml"'
    return resp


@app.route("/generate", methods=["POST"])
def generate():
    data = request.get_json(force=True)
    if not data:
        return jsonify({"error": "No JSON body provided"}), 400

    try:
        xml = generate_jats(data)
        return jsonify({"xml": xml})
    except Exception as e:
        return jsonify({"error": f"Generation failed: {str(e)}"}), 500


@app.route("/validate", methods=["POST"])
def validate():
    data = request.get_json(force=True)
    if not data or "xml" not in data:
        return jsonify({"error": "Expected JSON body with 'xml' key"}), 400

    try:
        result = validate_jats(data["xml"])
        return jsonify(result)
    except Exception as e:
        return jsonify({"error": f"Validation failed: {str(e)}"}), 500


@app.route("/autotag", methods=["POST"])
def autotag():
    api_key = os.getenv("ANTHROPIC_API_KEY")
    if not api_key:
        return jsonify({"error": "ANTHROPIC_API_KEY not configured on server"}), 503

    data = request.get_json(force=True)
    if not data or "text" not in data:
        return jsonify({"error": "Expected JSON body with 'text' key"}), 400

    raw_text = data["text"][:12000]  # Limit input to keep costs low

    try:
        client = anthropic.Anthropic(api_key=api_key)

        system_prompt = (
            "You are a JATS XML tagging assistant for academic journals. "
            "Given raw article text, identify and return JSON with: "
            "article_type (one of: Research Article, Review, Conference Proceeding, "
            "Enhanced Poster Abstract, Conference Report), "
            "sections (array of {heading: string, type: one of Introduction, Methods, "
            "Results, Discussion, Conclusion, Acknowledgements, Other}), "
            "and missing_sections (array of section type strings that are required but absent). "
            "Return only valid JSON with no markdown fences."
        )

        message = client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=1024,
            system=system_prompt,
            messages=[{"role": "user", "content": raw_text}],
        )

        response_text = message.content[0].text.strip()
        # Strip markdown fences if model adds them despite instructions
        response_text = response_text.removeprefix("```json").removeprefix("```").removesuffix("```").strip()

        parsed = json.loads(response_text)
        return jsonify(parsed)
    except json.JSONDecodeError as e:
        return jsonify({"error": f"AI returned invalid JSON: {e}"}), 500
    except Exception as e:
        return jsonify({"error": f"AI auto-tag failed: {str(e)}"}), 500


@app.route("/preview/web", methods=["POST"])
def preview_web():
    data = request.get_json(force=True)
    if not data:
        return jsonify({"error": "No JSON body provided"}), 400
    try:
        from pdf_gen.html_web_template import build_web_html
        html_str = build_web_html(data)
        resp = make_response(html_str.encode("utf-8"))
        resp.headers["Content-Type"] = "text/html; charset=utf-8"
        return resp
    except Exception as e:
        return jsonify({"error": f"Web preview failed: {str(e)}"}), 500


@app.route("/export/web-zip", methods=["POST"])
def export_web_zip():
    import zipfile, io as _io
    data = request.get_json(force=True)
    if not data:
        return jsonify({"error": "No JSON body provided"}), 400
    try:
        from pdf_gen.html_web_template import build_web_html
        html_str = build_web_html(data)
        slug = re.sub(r"[^a-zA-Z0-9_\-]", "_", data.get("title", "article"))[:60].strip("_") or "article"
        buf = _io.BytesIO()
        with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
            zf.writestr(f"{slug}/index.html", html_str.encode("utf-8"))
        buf.seek(0)
        resp = make_response(buf.read())
        resp.headers["Content-Type"] = "application/zip"
        resp.headers["Content-Disposition"] = f'attachment; filename="{slug}.zip"'
        return resp
    except Exception as e:
        return jsonify({"error": f"Web ZIP generation failed: {str(e)}"}), 500


@app.route("/export/xml-zip", methods=["POST"])
def export_xml_zip():
    """
    Return a ZIP package containing:
      - article.xml  (JATS XML with filename-based graphic hrefs)
      - figure images extracted from the parsed data_uri fields
        (named using the same convention as the XML: Journal-vol-issue-gN.jpeg)

    Upload both to the journal website so the XML graphic hrefs resolve.
    """
    import zipfile, io as _io, base64
    data = request.get_json(force=True)
    if not data:
        return jsonify({"error": "No JSON body provided"}), 400
    try:
        from jats.generator import generate_jats, fig_filename, eq_filename

        xml_str = generate_jats(data)
        slug    = re.sub(r"[^a-zA-Z0-9_\-]", "_", data.get("title", "article"))[:60].strip("_") or "article"

        buf = _io.BytesIO()
        fig_counter = 0
        eq_counter  = 0
        with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
            zf.writestr("article.xml", xml_str.encode("utf-8"))

            # Walk sections content in document order — extract figures and equations
            for sec in data.get("sections", []):
                all_content = [sec.get("content", [])] + \
                              [s.get("content", []) for s in sec.get("subsections", [])]
                for content_list in all_content:
                    for block in content_list:
                        btype = block.get("type")
                        uri   = block.get("data_uri", "")
                        if not uri:
                            if btype in ("figure", "equation"):
                                if btype == "figure":  fig_counter += 1
                                else:                  eq_counter  += 1
                            continue

                        b64 = uri.split(",", 1)[-1] if "," in uri else uri
                        try:
                            img_bytes = base64.b64decode(b64)
                        except Exception:
                            img_bytes = None

                        if btype == "figure":
                            fig_counter += 1
                            fname = fig_filename(data, fig_counter)
                            if img_bytes:
                                zf.writestr(fname, img_bytes)
                        elif btype == "equation":
                            eq_counter += 1
                            fname = eq_filename(data, eq_counter)
                            if img_bytes:
                                zf.writestr(fname, img_bytes)

        buf.seek(0)
        resp = make_response(buf.read())
        resp.headers["Content-Type"] = "application/zip"
        resp.headers["Content-Disposition"] = f'attachment; filename="{slug}_xml_package.zip"'
        return resp
    except Exception as e:
        return jsonify({"error": f"XML package generation failed: {str(e)}"}), 500


@app.route("/export/html", methods=["POST"])
def export_html():
    data = request.get_json(force=True)
    if not data:
        return jsonify({"error": "No JSON body provided"}), 400

    try:
        from pdf_gen.html_template import build_html

        html_str = build_html(data)
        slug = re.sub(r"[^a-zA-Z0-9_\-]", "_", data.get("title", "article"))[:60].strip("_") or "article"
        resp = make_response(html_str.encode("utf-8"))
        resp.headers["Content-Type"] = "text/html; charset=utf-8"
        resp.headers["Content-Disposition"] = f'attachment; filename="{slug}.html"'
        return resp
    except Exception as e:
        return jsonify({"error": f"HTML generation failed: {str(e)}"}), 500


@app.route("/export/pdf", methods=["POST"])
def export_pdf():
    data = request.get_json(force=True)
    if not data:
        return jsonify({"error": "No JSON body provided"}), 400

    try:
        from pdf_gen.html_template import build_html
        from pdf_gen.renderer import render_pdf

        # Enhanced Poster Abstract uses single-column layout
        use_two_col = data.get("article_type") != "Enhanced Poster Abstract"
        html_str  = build_html(data, two_col=use_two_col)
        pdf_bytes = render_pdf(html_str)

        slug = re.sub(r"[^a-zA-Z0-9_\-]", "_", data.get("title", "article"))[:60].strip("_") or "article"
        resp = make_response(pdf_bytes)
        resp.headers["Content-Type"] = "application/pdf"
        resp.headers["Content-Disposition"] = f'attachment; filename="{slug}.pdf"'
        return resp
    except Exception as e:
        return jsonify({"error": f"PDF generation failed: {str(e)}"}), 500


@app.route("/export/word", methods=["POST"])
def export_word():
    """Export article as Microsoft Word (.docx) file with full formatting"""
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.shared import OxmlElement
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls, qn
    import base64
    from io import BytesIO
    from PIL import Image

    data = request.get_json(force=True)
    if not data:
        return jsonify({"error": "No JSON body provided"}), 400

    try:
        doc = Document()

        # Header with journal info and logos
        journal_name = data.get("journal_name", "Novel Future Publishing")
        publisher_name = data.get("publisher_name", "")
        journal_logo = data.get("journal_logo", "")  # Logo image from metadata
        publisher_logo = data.get("publisher_logo", "")  # Publisher logo from metadata

        header_table = doc.add_table(rows=1, cols=3)
        header_table.autofit = False
        header_table.allow_autofit = False

        # Set column widths
        for cell in header_table.rows[0].cells:
            cell.width = Inches(1.8)

        # Left cell - Journal logo (or initials fallback)
        left_cell = header_table.rows[0].cells[0]
        left_cell.paragraphs[0].clear()
        left_para = left_cell.paragraphs[0]
        left_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if journal_logo and journal_logo.startswith("data:image"):
            try:
                # Extract and insert logo image
                b64_data = journal_logo.split(",")[1]
                logo_data = base64.b64decode(b64_data)
                logo_stream = BytesIO(logo_data)
                left_para.add_run().add_picture(logo_stream, width=Inches(1.2))
            except Exception as e:
                print(f"⚠ Journal logo insert failed: {e}, using initials")
                initials = "".join(w[0].upper() for w in journal_name.split()[:2]) or "NP"
                left_run = left_para.add_run(initials)
                left_run.font.size = Pt(32)
                left_run.font.bold = True
                left_run.font.color.rgb = RGBColor(15, 53, 87)
        else:
            # Fallback to initials
            initials = "".join(w[0].upper() for w in journal_name.split()[:2]) or "NP"
            left_run = left_para.add_run(initials)
            left_run.font.size = Pt(32)
            left_run.font.bold = True
            left_run.font.color.rgb = RGBColor(15, 53, 87)

        # Center cell - Journal name
        center_cell = header_table.rows[0].cells[1]
        center_cell.paragraphs[0].clear()
        center_para = center_cell.paragraphs[0]
        center_run = center_para.add_run("From the journal:\n")
        center_run.font.size = Pt(8)
        center_run.font.color.rgb = RGBColor(128, 128, 128)
        journal_run = center_para.add_run(journal_name)
        journal_run.font.size = Pt(14)
        journal_run.font.bold = True
        journal_run.font.color.rgb = RGBColor(15, 53, 87)
        center_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Right cell - Publisher logo (or name fallback)
        right_cell = header_table.rows[0].cells[2]
        right_cell.paragraphs[0].clear()
        right_para = right_cell.paragraphs[0]
        right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        if publisher_logo and publisher_logo.startswith("data:image"):
            try:
                # Extract and insert publisher logo image
                b64_data = publisher_logo.split(",")[1]
                logo_data = base64.b64decode(b64_data)
                logo_stream = BytesIO(logo_data)
                right_para.add_run().add_picture(logo_stream, width=Inches(1.5))
            except Exception as e:
                print(f"⚠ Publisher logo insert failed: {e}, using text")
                if publisher_name:
                    right_run = right_para.add_run(publisher_name)
                    right_run.font.size = Pt(10)
                    right_run.font.bold = True
        else:
            # Fallback to text
            if publisher_name:
                right_run = right_para.add_run(publisher_name)
                right_run.font.size = Pt(10)
                right_run.font.bold = True

        # Remove table borders for cleaner look
        tbl = header_table._element
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)
        tblBorders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'none')
            border.set(qn('w:sz'), '0')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), 'auto')
            tblBorders.append(border)
        tblPr.append(tblBorders)

        # Add horizontal rule
        hr_para = doc.add_paragraph()
        hr_para.paragraph_format.space_before = Pt(6)
        hr_para.paragraph_format.space_after = Pt(12)
        pPr = hr_para._element.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '24')  # 3pt
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '0F3557')
        pBdr.append(bottom)
        pPr.append(pBdr)

        # Article type badge
        article_type = data.get("article_type", "Research Article")
        badge_para = doc.add_paragraph()
        badge_run = badge_para.add_run(article_type.upper())
        badge_run.font.size = Pt(9)
        badge_run.font.bold = True
        badge_run.font.color.rgb = RGBColor(255, 255, 255)
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), '0F3557')
        badge_para._element.get_or_add_pPr().append(shading_elm)
        badge_para.paragraph_format.left_indent = Inches(0)
        badge_para.paragraph_format.right_indent = Inches(0)
        badge_para.paragraph_format.space_before = Pt(6)
        badge_para.paragraph_format.space_after = Pt(12)

        # Title
        title = data.get("title", "Untitled")
        title_para = doc.add_paragraph(title, style="Heading 1")
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_para.paragraph_format.space_after = Pt(12)
        title_para.paragraph_format.line_spacing = 1.5

        # Authors with affiliations
        authors = data.get("authors", [])
        if authors:
            # Author line with superscript affiliation numbers
            author_names = []
            for a in authors:
                name_parts = [a.get("first_name", ""), a.get("last_name", "")]
                name = " ".join([p for p in name_parts if p]).strip()
                # Add superscript affiliation marker if available
                affil = a.get("affiliation", "")
                if name:
                    author_names.append(name)

            if author_names:
                authors_para = doc.add_paragraph(", ".join(author_names))
                authors_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                authors_para.paragraph_format.space_after = Pt(6)
                authors_para.paragraph_format.line_spacing = 1.15

        # Collect unique affiliations from authors
        unique_affiliations = {}
        for a in authors:
            aff = a.get("affiliation", "").strip()
            if aff and aff not in unique_affiliations:
                unique_affiliations[aff] = len(unique_affiliations) + 1

        # Numbered affiliations
        if unique_affiliations:
            for aff_text, aff_num in sorted(unique_affiliations.items(), key=lambda x: x[1]):
                aff_para = doc.add_paragraph()
                aff_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                # Add superscript number
                run = aff_para.add_run(f"{aff_num}")
                run.font.superscript = True
                aff_para.add_run(aff_text)

        # Corresponding author and email
        corresp_author = None
        corresp_emails = []
        for a in authors:
            if a.get("corresponding"):
                name_parts = [a.get("first_name", ""), a.get("last_name", "")]
                name = " ".join([p for p in name_parts if p]).strip()
                if name:
                    corresp_author = name
                email = a.get("email", "")
                if email:
                    corresp_emails.append(email)

        if corresp_author or corresp_emails:
            corresp_para = doc.add_paragraph()
            corresp_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            corresp_run = corresp_para.add_run("Corresponding author: ")
            corresp_run.bold = True
            if corresp_author:
                corresp_para.add_run(corresp_author)
            if corresp_emails:
                corresp_para.add_run(f" ({', '.join(corresp_emails)})")

        doc.add_paragraph()  # Spacing

        # Abstract
        if data.get("abstract"):
            abstract_heading = doc.add_heading("Abstract", level=2)
            abstract_heading.paragraph_format.space_before = Pt(12)
            abstract_heading.paragraph_format.space_after = Pt(6)
            abstract_para = doc.add_paragraph(data["abstract"])
            abstract_para.paragraph_format.space_after = Pt(12)
            abstract_para.paragraph_format.line_spacing = 1.5
            doc.add_paragraph()  # Spacing

        def add_content_block(block):
            """Helper to add a single content block to the document."""
            if not isinstance(block, dict):
                return

            block_type = block.get("type")

            # Paragraph block
            if block_type == "paragraph":
                text = block.get("text", "")
                if text:
                    para = doc.add_paragraph(text)
                    para.paragraph_format.line_spacing = 1.5
                    para.paragraph_format.space_after = Pt(6)

            # Figure/Image block
            elif block_type == "figure":
                data_uri = block.get("data_uri", "")
                if data_uri and data_uri.startswith("data:image"):
                    try:
                        b64_data = data_uri.split(",")[1]
                        img_data = base64.b64decode(b64_data)
                        img_stream = BytesIO(img_data)
                        doc.add_picture(img_stream, width=Inches(5.5))
                    except Exception:
                        pass
                label = block.get("label", "Figure")
                caption = block.get("caption", "")
                if label or caption:
                    fig_text = f"{label}. {caption}" if caption else label
                    fig_para = doc.add_paragraph(fig_text)
                    if fig_para.runs:
                        fig_para.runs[0].italic = True
                    fig_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Table block
            elif block_type == "table":
                headers = block.get("headers", [])
                rows = block.get("rows", [])
                if headers or rows:
                    num_cols = len(headers) if headers else (len(rows[0]) if rows else 1)
                    num_rows = len(rows) + (1 if headers else 0)
                    table = doc.add_table(rows=num_rows, cols=num_cols)
                    table.style = "Light Grid Accent 1"
                    if headers:
                        for i, h in enumerate(headers):
                            if i < len(table.rows[0].cells):
                                table.rows[0].cells[i].text = str(h or "")
                    for r_idx, row in enumerate(rows):
                        row_offset = 1 if headers else 0
                        if r_idx + row_offset < len(table.rows):
                            for c_idx, cell in enumerate(row):
                                if c_idx < len(table.rows[r_idx + row_offset].cells):
                                    table.rows[r_idx + row_offset].cells[c_idx].text = str(cell or "")

            # Equation block - Image first, then text fallback
            elif block_type == "equation":
                omml = block.get("omml", "")
                eq_text = block.get("text", "")
                data_uri = block.get("data_uri", "")
                eq_added = False

                # Try image first (most reliable)
                if data_uri and data_uri.startswith("data:image"):
                    try:
                        b64_data = data_uri.split(",")[1]
                        img_data = base64.b64decode(b64_data)
                        img_stream = BytesIO(img_data)
                        eq_para = doc.add_paragraph()
                        eq_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        eq_para.paragraph_format.space_before = Pt(6)
                        eq_para.paragraph_format.space_after = Pt(6)
                        run = eq_para.add_run()
                        run.add_picture(img_stream, width=Inches(4.0))
                        eq_added = True
                        print(f"✓ Equation image inserted successfully")
                    except Exception as e:
                        print(f"⚠ Equation image insertion failed: {e}")

                # Try OMML if image failed
                if not eq_added and omml:
                    try:
                        p = doc.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p.paragraph_format.space_before = Pt(6)
                        p.paragraph_format.space_after = Pt(6)
                        r = p.add_run()
                        omml_elem = parse_xml(f'<w:r {nsdecls("w", "m")}>{omml}</w:r>')
                        r._element.getparent().replace(r._element, omml_elem)
                        eq_added = True
                        print(f"✓ Equation OMML inserted successfully")
                    except Exception as e:
                        print(f"⚠ OMML insertion failed: {e}")

                # Fallback to text
                if not eq_added and eq_text:
                    try:
                        eq_para = doc.add_paragraph(eq_text)
                        eq_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        eq_para.paragraph_format.space_before = Pt(6)
                        eq_para.paragraph_format.space_after = Pt(6)
                        eq_added = True
                        print(f"✓ Equation text fallback used")
                    except Exception as e:
                        print(f"⚠ Equation text insertion failed: {e}")

        # Sections with formatted content
        sections = data.get("sections", [])
        for sec in sections:
            sec_title = sec.get("heading") or sec.get("title") if isinstance(sec, dict) else sec
            if sec_title:
                sec_heading = doc.add_heading(str(sec_title), level=2)
                sec_heading.paragraph_format.space_before = Pt(12)
                sec_heading.paragraph_format.space_after = Pt(6)

            # Handle top-level content blocks
            content = sec.get("content") if isinstance(sec, dict) else None
            if content and isinstance(content, list):
                for block in content:
                    add_content_block(block)

            # Handle subsections
            subsections = sec.get("subsections", []) if isinstance(sec, dict) else []
            for subsec in subsections:
                subsec_title = subsec.get("heading", "") if isinstance(subsec, dict) else ""
                if subsec_title:
                    subsec_heading = doc.add_heading(str(subsec_title), level=3)
                    subsec_heading.paragraph_format.space_before = Pt(12)
                    subsec_heading.paragraph_format.space_after = Pt(6)

                subsec_content = subsec.get("content", []) if isinstance(subsec, dict) else []
                if subsec_content and isinstance(subsec_content, list):
                    for block in subsec_content:
                        add_content_block(block)

        doc.add_paragraph()  # Spacing

        # References
        references = data.get("references", [])
        if references:
            doc.add_paragraph()  # Spacing before references
            ref_heading = doc.add_heading("References", level=2)
            ref_heading.paragraph_format.space_before = Pt(12)
            ref_heading.paragraph_format.space_after = Pt(6)
            ref_count = 0

            for ref in references:
                # Extract reference text - handle multiple formats
                ref_text = ""
                if isinstance(ref, dict):
                    # Parser format: {"number": i, "raw_text": text, "doi": doi}
                    ref_text = ref.get("raw_text") or ref.get("text", "")
                elif isinstance(ref, str):
                    ref_text = ref.strip()
                else:
                    ref_text = str(ref).strip() if ref else ""

                # Skip empty references
                if not ref_text:
                    continue

                ref_count += 1
                # Add numbered reference paragraph with proper indentation
                ref_para = doc.add_paragraph()
                ref_para.paragraph_format.left_indent = Inches(0.5)
                ref_para.paragraph_format.first_line_indent = Inches(-0.5)
                ref_para.paragraph_format.space_after = Pt(6)
                ref_para.paragraph_format.line_spacing = 1.15

                # Add number and text
                num_run = ref_para.add_run(f"{ref_count}. ")
                num_run.bold = True
                ref_para.add_run(ref_text)

            if ref_count == 0:
                # No valid references found
                ref_para = doc.add_paragraph("No references provided")
                ref_para.runs[0].italic = True
        else:
            print("⚠ WARNING: No references found in parsed data!")

        # Generate file
        slug = re.sub(r"[^a-zA-Z0-9_\-]", "_", title)[:60].strip("_") or "article"
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            doc.save(tmp.name)
            with open(tmp.name, "rb") as f:
                docx_bytes = f.read()
            os.unlink(tmp.name)

        resp = make_response(docx_bytes)
        resp.headers["Content-Type"] = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        resp.headers["Content-Disposition"] = f'attachment; filename="{slug}.docx"'
        return resp
    except Exception as e:
        return jsonify({"error": f"Word export failed: {str(e)}"}), 500


def _crossref_item_to_vancouver(item: dict) -> str:
    """
    Format a Crossref work metadata dict as a Vancouver-style citation.
    Pattern: Authors. Title. Journal. Year;Vol(Issue):Pages. doi:DOI
    """
    # ── Authors (max 6, then et al.) ────────────────────────────────────────
    authors = item.get("author", [])
    author_parts = []
    for a in authors[:6]:
        family  = a.get("family", "").strip()
        given   = a.get("given",  "").strip()
        initials = "".join(w[0].upper() for w in given.split() if w)
        author_parts.append(f"{family} {initials}".strip() if family else "")
    author_parts = [p for p in author_parts if p]
    author_str = ", ".join(author_parts)
    if len(authors) > 6:
        author_str += " et al"

    # ── Title ────────────────────────────────────────────────────────────────
    titles = item.get("title", [])
    title  = titles[0].strip() if titles else ""

    # ── Journal (prefer abbreviated title) ───────────────────────────────────
    short = item.get("short-container-title", [])
    journal = (short[0] if short else
               (item.get("container-title", [""])[0])).strip()

    # ── Publication date ─────────────────────────────────────────────────────
    for date_key in ("published", "published-print", "published-online"):
        dp = item.get(date_key, {}).get("date-parts", [[""]])
        if dp and dp[0] and dp[0][0]:
            year = str(dp[0][0])
            break
    else:
        year = ""

    volume = item.get("volume", "")
    issue  = item.get("issue",  "")
    pages  = item.get("page",   "")
    doi    = item.get("DOI",    "").strip()

    # ── Assemble ─────────────────────────────────────────────────────────────
    parts = []
    if author_str:
        parts.append(author_str + ".")
    if title:
        parts.append(title + ".")
    if journal:
        loc = year
        if volume:
            loc += f";{volume}"
        if issue:
            loc += f"({issue})"
        if pages:
            loc += f":{pages}"
        parts.append(f"{journal}. {loc}.".strip())
    if doi:
        parts.append(f"https://doi.org/{doi}")

    return " ".join(parts)


@app.route("/export/pdf-latex", methods=["POST"])
def export_pdf_latex():
    """Export article as PDF using LaTeX for professional mathematical typesetting"""
    data = request.get_json(force=True)
    if not data:
        return jsonify({"error": "No JSON body provided"}), 400

    try:
        from utils.pdf_latex import generate_pdf_from_latex

        pdf_bytes = generate_pdf_from_latex(data)

        slug = re.sub(r"[^a-zA-Z0-9_\-]", "_", data.get("title", "article"))[:60].strip("_") or "article"
        resp = make_response(pdf_bytes)
        resp.headers["Content-Type"] = "application/pdf"
        resp.headers["Content-Disposition"] = f'attachment; filename="{slug}_latex.pdf"'
        return resp
    except Exception as e:
        return jsonify({"error": f"LaTeX PDF generation failed: {str(e)}"}), 500


@app.route("/export/poster-web", methods=["POST"])
def export_poster_web():
    """Export poster as responsive HTML preview"""
    data = request.get_json(force=True)
    if not data:
        return jsonify({"error": "No JSON body provided"}), 400

    try:
        from utils.poster_web import generate_poster_html

        html = generate_poster_html(data)
        resp = make_response(html)
        resp.headers["Content-Type"] = "text/html; charset=utf-8"
        return resp
    except Exception as e:
        return jsonify({"error": f"Poster HTML generation failed: {str(e)}"}), 500


@app.route("/export/poster-word", methods=["POST"])
def export_poster_word():
    """Export poster as Word document (metadata + image)"""
    data = request.get_json(force=True)
    if not data:
        return jsonify({"error": "No JSON body provided"}), 400

    try:
        from utils.poster_word import generate_poster_word

        word_bytes = generate_poster_word(data)
        slug = re.sub(r"[^a-zA-Z0-9_\-]", "_", data.get("title", "poster"))[:60].strip("_") or "poster"
        resp = make_response(word_bytes)
        resp.headers["Content-Type"] = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        resp.headers["Content-Disposition"] = f'attachment; filename="{slug}.docx"'
        return resp
    except Exception as e:
        return jsonify({"error": f"Poster Word generation failed: {str(e)}"}), 500


@app.route("/export/poster-pdf-latex", methods=["POST"])
def export_poster_pdf_latex():
    """Export poster as PDF using LaTeX (metadata + image on separate pages)"""
    data = request.get_json(force=True)
    if not data:
        return jsonify({"error": "No JSON body provided"}), 400

    try:
        from utils.poster_pdf_latex import generate_poster_pdf_latex

        pdf_bytes = generate_poster_pdf_latex(data)
        slug = re.sub(r"[^a-zA-Z0-9_\-]", "_", data.get("title", "poster"))[:60].strip("_") or "poster"
        resp = make_response(pdf_bytes)
        resp.headers["Content-Type"] = "application/pdf"
        resp.headers["Content-Disposition"] = f'attachment; filename="{slug}.pdf"'
        return resp
    except Exception as e:
        return jsonify({"error": f"Poster PDF generation failed: {str(e)}"}), 500


@app.route("/export/poster-html", methods=["POST"])
def export_poster_html():
    """Export poster as HTML document"""
    data = request.get_json(force=True)
    if not data:
        return jsonify({"error": "No JSON body provided"}), 400

    try:
        from utils.poster_html import generate_poster_html_document

        html = generate_poster_html_document(data)
        slug = re.sub(r"[^a-zA-Z0-9_\-]", "_", data.get("title", "poster"))[:60].strip("_") or "poster"
        resp = make_response(html)
        resp.headers["Content-Type"] = "text/html; charset=utf-8"
        resp.headers["Content-Disposition"] = f'attachment; filename="{slug}.html"'
        return resp
    except Exception as e:
        return jsonify({"error": f"Poster HTML generation failed: {str(e)}"}), 500


@app.route("/export/poster-pdf", methods=["POST"])
def export_poster_pdf_json():
    """Export poster as PDF from JSON data (parsed poster metadata)"""
    data = request.get_json(force=True)
    if not data:
        return jsonify({"error": "No JSON body provided"}), 400

    # Debug logging
    logger.info(f"[POSTER-PDF] Received data fields: {list(data.keys())}")
    poster_image = data.get("poster_image", "")
    logger.info(f"[POSTER-PDF] poster_image size: {len(poster_image)} bytes")
    logger.info(f"[POSTER-PDF] Title: {data.get('title')}")
    logger.info(f"[POSTER-PDF] Authors: {len(data.get('authors', []))} author(s)")

    # Skip large images (> 10MB base64) to avoid memory issues
    if len(poster_image) > 10 * 1024 * 1024:
        logger.warning(f"[POSTER-PDF] Image too large ({len(poster_image)} bytes), skipping")
        data["poster_image"] = ""

    try:
        from utils.poster_pdf_json import generate_poster_pdf_from_json

        journal_logo = data.get("journal_logo", "")
        publisher_logo = data.get("publisher_logo", "")

        pdf_bytes = generate_poster_pdf_from_json(data, journal_logo, publisher_logo)
        logger.info(f"[POSTER-PDF] Generated PDF: {len(pdf_bytes)} bytes")

        slug = re.sub(r"[^a-zA-Z0-9_\-]", "_", data.get("title", "poster"))[:60].strip("_") or "poster"
        resp = make_response(pdf_bytes)
        resp.headers["Content-Type"] = "application/pdf"
        resp.headers["Content-Disposition"] = f'attachment; filename="{slug}.pdf"'
        return resp
    except Exception as e:
        return jsonify({"error": f"Poster PDF generation failed: {str(e)}"}), 500


@app.route("/export/poster-pdf-docx", methods=["POST"])
def export_poster_pdf_docx():
    """Export poster as PDF using LibreOffice (converts DOCX with embedded images and logos)"""
    if "file" not in request.files:
        return jsonify({"error": "No file provided"}), 400

    file = request.files["file"]
    if file.filename == "" or not allowed_file(file.filename):
        return jsonify({"error": "Invalid file"}), 400

    # Optional: get logos from form data
    journal_logo = request.form.get("journal_logo", "")
    publisher_logo = request.form.get("publisher_logo", "")

    try:
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            file.save(tmp.name)
            tmp_path = tmp.name

        from utils.poster_pdf_libreoffice import generate_poster_pdf_libreoffice
        from utils.pdf_logos import add_logos_to_pdf
        from parser.poster_parser import parse_poster

        # Convert DOCX to PDF
        try:
            pdf_bytes = generate_poster_pdf_libreoffice(tmp_path)
        except Exception as e:
            # Log the full error for debugging
            import traceback
            print(f"ERROR in PDF generation: {str(e)}")
            print(f"Traceback: {traceback.format_exc()}")
            raise

        # Extract poster metadata and image from DOCX
        try:
            print(f"[POSTER] Parsing DOCX for image extraction...")
            poster_data = parse_poster(tmp_path)
            poster_image = poster_data.get("poster_image", "")

            print(f"[POSTER] Image extracted: {len(poster_image)} bytes")
            print(f"[POSTER] Image present: {bool(poster_image)}")

            # Add poster image as overlay on first page
            if poster_image:
                print(f"[POSTER] Adding image overlay to PDF...")
                from utils.pdf_logos import add_logos_to_pdf
                pdf_before = len(pdf_bytes)
                # Use the image embedding system to add poster as "publisher logo"
                pdf_bytes = add_logos_to_pdf(pdf_bytes, "", poster_image)
                pdf_after = len(pdf_bytes)
                print(f"[POSTER] PDF size before: {pdf_before}, after: {pdf_after}")
            else:
                print(f"[POSTER] No image found in poster data")
        except Exception as e:
            import traceback
            print(f"[POSTER] Error embedding image: {e}")
            print(f"[POSTER] Traceback: {traceback.format_exc()}")

        # Add logos if provided
        if journal_logo or publisher_logo:
            pdf_bytes = add_logos_to_pdf(pdf_bytes, journal_logo, publisher_logo)

        slug = re.sub(r"[^a-zA-Z0-9_\-]", "_", file.filename.rsplit(".", 1)[0])[:60].strip("_") or "poster"
        resp = make_response(pdf_bytes)
        resp.headers["Content-Type"] = "application/pdf"
        resp.headers["Content-Disposition"] = f'attachment; filename="{slug}.pdf"'
        return resp
    except Exception as e:
        return jsonify({"error": f"Poster PDF generation failed: {str(e)}"}), 500
    finally:
        try:
            os.unlink(tmp_path)
        except Exception:
            pass


@app.route("/export/poster-xml", methods=["POST"])
def export_poster_xml():
    """Export poster as JATS XML"""
    data = request.get_json(force=True)
    if not data:
        return jsonify({"error": "No JSON body provided"}), 400

    try:
        from utils.poster_xml import generate_poster_xml

        xml = generate_poster_xml(data)
        slug = re.sub(r"[^a-zA-Z0-9_\-]", "_", data.get("title", "poster"))[:60].strip("_") or "poster"
        resp = make_response(xml)
        resp.headers["Content-Type"] = "application/xml; charset=utf-8"
        resp.headers["Content-Disposition"] = f'attachment; filename="{slug}.xml"'
        return resp
    except Exception as e:
        return jsonify({"error": f"Poster XML generation failed: {str(e)}"}), 500


@app.route("/format-refs", methods=["POST"])
def format_refs():
    """
    Reformat all references into Vancouver citation style using Crossref
    metadata. For refs with a DOI, metadata is fetched directly. For refs
    without a DOI, a bibliographic search is attempted first.
    Returns the same refs list with raw_text set to the Vancouver string
    and doi filled in where found.
    """
    import requests as _req
    import time

    data = request.get_json(force=True)
    if not data or "refs" not in data:
        return jsonify({"error": "Expected JSON body with 'refs' key"}), 400

    refs    = data["refs"]
    enriched = []
    formatted = 0

    for ref in refs:
        raw_text = ref.get("raw_text", "").strip()
        doi      = ref.get("doi", "").strip()
        item     = None

        try:
            if doi:
                # Fetch metadata directly by DOI
                r = _req.get(
                    f"https://api.crossref.org/works/{doi}",
                    timeout=6,
                    headers={"User-Agent": "NovelTOR/1.0 (mailto:support@noveltor.com)"},
                )
                if r.status_code == 200:
                    item = r.json().get("message", {})
            else:
                # Bibliographic search
                r = _req.get(
                    "https://api.crossref.org/works",
                    params={"query.bibliographic": raw_text, "rows": 1,
                            "select": "DOI,score,title,author,container-title,"
                                      "short-container-title,published,published-print,"
                                      "published-online,volume,issue,page"},
                    timeout=6,
                    headers={"User-Agent": "NovelTOR/1.0 (mailto:support@noveltor.com)"},
                )
                items = r.json().get("message", {}).get("items", [])
                if items and items[0].get("score", 0) > 50:
                    item = items[0]
                    doi  = item.get("DOI", "").strip()

            if item:
                vancouver = _crossref_item_to_vancouver(item)
                if vancouver:
                    new_ref = dict(ref)
                    new_ref["raw_text"] = vancouver
                    new_ref["doi"]      = doi
                    enriched.append(new_ref)
                    formatted += 1
                    time.sleep(0.12)
                    continue

        except Exception:
            pass  # fall through to keep original

        enriched.append(ref)
        time.sleep(0.12)

    return jsonify({"refs": enriched, "formatted": formatted})


@app.route("/enrich-refs", methods=["POST"])
def enrich_refs():
    """
    For each reference that has no DOI, query the Crossref API using the
    raw_text as a bibliographic search string. Returns the same refs list
    with doi fields filled in where a confident match was found.
    """
    import requests as _req
    import time

    data = request.get_json(force=True)
    if not data or "refs" not in data:
        return jsonify({"error": "Expected JSON body with 'refs' key"}), 400

    refs = data["refs"]
    enriched = []
    found = 0

    for ref in refs:
        raw_text = ref.get("raw_text", "").strip()
        existing_doi = ref.get("doi", "").strip()

        # Already has DOI — skip Crossref lookup
        if existing_doi:
            enriched.append(ref)
            continue

        if not raw_text:
            enriched.append(ref)
            continue

        try:
            resp = _req.get(
                "https://api.crossref.org/works",
                params={"query.bibliographic": raw_text, "rows": 1, "select": "DOI,score,title"},
                timeout=6,
                headers={"User-Agent": "NovelTOR/1.0 (mailto:support@noveltor.com)"},
            )
            items = resp.json().get("message", {}).get("items", [])
            doi = ""
            if items:
                top = items[0]
                score = top.get("score", 0)
                # Crossref scores > 50 indicate a confident bibliographic match
                if score > 50:
                    doi = top.get("DOI", "").strip()

            new_ref = dict(ref)
            new_ref["doi"] = doi
            if doi:
                found += 1
            enriched.append(new_ref)

            # Be polite to Crossref — 10 req/s max
            time.sleep(0.12)

        except Exception:
            enriched.append(ref)

    return jsonify({"refs": enriched, "found": found})


if __name__ == "__main__":
    port = int(os.getenv("FLASK_PORT", 5000))
    debug = os.getenv("FLASK_DEBUG", "true").lower() == "true"
    app.run(host="0.0.0.0", port=port, debug=debug)
